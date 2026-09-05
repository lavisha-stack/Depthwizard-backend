from __future__ import annotations

import hashlib
import math
import shutil
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK, WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\preet\OneDrive\Documents\ChatGPT\Hackathon 2026 Depth wizard")
REFERENCE = Path(
    r"C:\Users\preet\.codex\plugins\cache\openai-curated-remote\openai-templates\0.1.1"
    r"\skills\artifact-template-experiment-analysis\assets\reference.docx"
)
REFERENCE_SHA256 = "D823CD0115186B34C01C6E4B4DA3BE28B64EE73CAC849DBD62D6F4BB6385B0FB"
TMP = ROOT / "docs" / "_tmp_depthwizard_guide"
ASSETS = TMP / "assets"
OUT_DIR = ROOT / "docs" / "output"
OUTPUT_DOCX = OUT_DIR / "DepthWizard_Team_Technical_and_Judge_Guide.docx"

GREEN = "416658"
DARK_GREEN = "29463C"
PALE_GREEN = "E9F1ED"
MID_GREEN = "AFC8BD"
LIGHT_GRAY = "F2F3F2"
MID_GRAY = "D8DDDA"
DARK = "202522"
MUTED = "5E6863"
WHITE = "FFFFFF"


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            h.update(block)
    return h.hexdigest().upper()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_borders(table, color=MID_GRAY, size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_paragraph_border(paragraph, color=GREEN, size="12", space="1") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_toc_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Update field to generate the table of contents."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, placeholder, end])


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Georgia"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    title = styles["Title"]
    title.font.name = "Georgia"
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(GREEN)
    title.paragraph_format.space_after = Pt(7)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Georgia"
    subtitle.font.size = Pt(14)
    subtitle.font.italic = False
    subtitle.font.color.rgb = RGBColor.from_string(MUTED)

    heading_specs = {
        "Heading 1": (22, GREEN, 14, 7),
        "Heading 2": (14, GREEN, 10, 4),
        "Heading 3": (10.5, DARK_GREEN, 7, 2),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = styles[name]
        style.font.name = "Georgia"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        if name not in [s.name for s in styles]:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = styles[name]
        style.font.name = "Georgia"
        style.font.size = Pt(9.3)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.left_indent = Inches(0.23)
        style.paragraph_format.first_line_indent = Inches(-0.14)

    if "DW Code" not in [s.name for s in styles]:
        style = styles.add_style("DW Code", 1)
    else:
        style = styles["DW Code"]
    style.font.name = "Consolas"
    style.font.size = Pt(7.8)
    style.font.color.rgb = RGBColor.from_string(DARK)
    style.paragraph_format.space_before = Pt(3)
    style.paragraph_format.space_after = Pt(5)
    style.paragraph_format.left_indent = Inches(0.18)
    style.paragraph_format.right_indent = Inches(0.18)

    if "DW Caption" not in [s.name for s in styles]:
        style = styles.add_style("DW Caption", 1)
    else:
        style = styles["DW Caption"]
    style.font.name = "Georgia"
    style.font.size = Pt(8)
    style.font.italic = True
    style.font.color.rgb = RGBColor.from_string(MUTED)
    style.paragraph_format.space_before = Pt(2)
    style.paragraph_format.space_after = Pt(7)
    style.paragraph_format.keep_with_next = False

    if "DW Eyebrow" not in [s.name for s in styles]:
        style = styles.add_style("DW Eyebrow", 1)
    else:
        style = styles["DW Eyebrow"]
    style.font.name = "Georgia"
    style.font.size = Pt(8)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(GREEN)
    style.paragraph_format.space_after = Pt(4)


def rebuild_footer(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        for child in list(footer._element):
            footer._element.remove(child)
        table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.columns[0].width = Inches(5.5)
        table.columns[1].width = Inches(1.0)
        table.rows[0].cells[0].width = Inches(5.5)
        table.rows[0].cells[1].width = Inches(1.0)
        left = table.cell(0, 0)
        right = table.cell(0, 1)
        for cell in (left, right):
            set_cell_margins(cell, 20, 0, 0, 0)
        lp = left.paragraphs[0]
        lp.text = "DEPTHWIZARD · TEAM TECHNICAL GUIDE"
        lp.style = doc.styles["DW Eyebrow"]
        rp = right.paragraphs[0]
        rp.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        run = rp.add_run("PAGE ")
        run.font.name = "Georgia"
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(GREEN)
        add_page_number_field(rp)


def clear_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def add_eyebrow(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="DW Eyebrow")
    p.add_run(text.upper())
    set_paragraph_border(p, color=MID_GREEN, size="5", space="5")


def add_h1(doc: Document, text: str, *, new_page: bool = True) -> None:
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.page_break_before = new_page
    p.add_run(text)
    set_paragraph_border(p, color=MID_GREEN, size="6", space="4")


def add_h2(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="Heading 2")


def add_h3(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="Heading 3")


def add_body(doc: Document, text: str, *, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        p.add_run(bold_lead).bold = True
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.24)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        p.paragraph_format.space_after = Pt(2)
        p.add_run("• ").font.color.rgb = RGBColor.from_string(GREEN)
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for index, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.30)
        p.paragraph_format.first_line_indent = Inches(-0.22)
        p.paragraph_format.space_after = Pt(2)
        number = p.add_run(f"{index}. ")
        number.bold = True
        number.font.color.rgb = RGBColor.from_string(GREEN)
        p.add_run(item)


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="DW Code")
    p.add_run(text)
    set_paragraph_shading(p, LIGHT_GRAY)


def add_note(doc: Document, title: str, text: str, *, warning: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.16)
    p.paragraph_format.right_indent = Inches(0.08)
    p.add_run(title.upper() + "  ").bold = True
    p.runs[0].font.color.rgb = RGBColor.from_string(GREEN if not warning else DARK_GREEN)
    p.add_run(text)
    set_paragraph_shading(p, PALE_GREEN if not warning else "F3EFE7")
    set_paragraph_border(p, color=GREEN, size="8", space="4")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None, font_size=8.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    if widths:
        for idx, width in enumerate(widths):
            table.columns[idx].width = Inches(width)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    prevent_row_split(hdr)
    for idx, value in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_shading(cell, GREEN)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(value)
        run.bold = True
        run.font.name = "Georgia"
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor.from_string(WHITE)
    for row_index, values in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        for idx, value in enumerate(values):
            cell = row.cells[idx]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if row_index % 2:
                set_cell_shading(cell, "F8F9F8")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(value))
            run.font.name = "Georgia"
            run.font.size = Pt(font_size)
            run.font.color.rgb = RGBColor.from_string(DARK)
        if widths:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_picture(doc: Document, path: Path, caption: str, width: float = 6.25) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption, style="DW Caption")
    cap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def add_two_pictures(doc: Document, left: Path, right: Path, left_caption: str, right_caption: str) -> None:
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(3.18)
    table.columns[1].width = Inches(3.18)
    for cell in table.rows[0].cells:
        set_cell_margins(cell, 0, 30, 30, 30)
    for idx, path in enumerate((left, right)):
        p = table.cell(0, idx).paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.add_run().add_picture(str(path), width=Inches(3.05))
    for idx, caption in enumerate((left_caption, right_caption)):
        cell = table.cell(1, idx)
        set_cell_margins(cell, 0, 45, 60, 45)
        p = cell.paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(caption)
        run.italic = True
        run.font.size = Pt(7.7)
        run.font.color.rgb = RGBColor.from_string(MUTED)


def add_page_break(doc: Document) -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _font(size: int, bold: bool = False):
    preferred = Path(r"C:\Windows\Fonts\georgiab.ttf" if bold else r"C:\Windows\Fonts\georgia.ttf")
    try:
        return ImageFont.truetype(str(preferred), size=size)
    except OSError:
        return ImageFont.load_default()


def _centre_text(draw, box, text, font, fill, spacing=4):
    left, top, right, bottom = box
    bbox = draw.multiline_textbbox((0, 0), text, font=font, align="center", spacing=spacing)
    width = bbox[2] - bbox[0]
    height = bbox[3] - bbox[1]
    draw.multiline_text(((left + right - width) / 2, (top + bottom - height) / 2), text, font=font, fill=fill, align="center", spacing=spacing)


def _box(draw, box, title, sub=None, fill=PALE_GREEN):
    draw.rounded_rectangle(box, radius=20, fill=f"#{fill}", outline=f"#{GREEN}", width=3)
    left, top, right, bottom = box
    if sub:
        _centre_text(draw, (left + 8, top + 9, right - 8, top + (bottom-top)*0.6), title, _font(25, True), f"#{DARK_GREEN}")
        _centre_text(draw, (left + 8, top + (bottom-top)*0.53, right - 8, bottom - 7), sub, _font(19), f"#{MUTED}")
    else:
        _centre_text(draw, box, title, _font(25, True), f"#{DARK_GREEN}")


def _arrow(draw, start, end, color=GREEN, width=5):
    draw.line([start, end], fill=f"#{color}", width=width)
    angle = math.atan2(end[1]-start[1], end[0]-start[0])
    length = 18
    points = [
        end,
        (end[0]-length*math.cos(angle-0.5), end[1]-length*math.sin(angle-0.5)),
        (end[0]-length*math.cos(angle+0.5), end[1]-length*math.sin(angle+0.5)),
    ]
    draw.polygon(points, fill=f"#{color}")


def _new_canvas(width=1800, height=760):
    image = Image.new("RGB", (width, height), "white")
    return image, ImageDraw.Draw(image)


def _save(image, name):
    path = ASSETS / name
    image.save(path, dpi=(180, 180), optimize=True)
    return path


def make_diagrams() -> dict[str, Path]:
    ASSETS.mkdir(parents=True, exist_ok=True)
    diagrams: dict[str, Path] = {}

    image, draw = _new_canvas(1900, 720)
    _centre_text(draw, (0, 15, 1900, 90), "DEPTHWIZARD · VERIFIED IMPLEMENTATION FLOW", _font(35, True), f"#{GREEN}")
    boxes = [
        (35, 180, 340, 355, "1 · INGEST", "RGB + metadata"),
        (415, 180, 720, 355, "2 · INFER", "relative depth"),
        (795, 180, 1100, 355, "3 · CALIBRATE", "SRTM / GCP"),
        (1175, 180, 1480, 355, "5 · ORCHESTRATE", "FastAPI jobs"),
        (1555, 180, 1860, 355, "6 · RENDER", "Three.js mesh"),
    ]
    for left, top, right, bottom, title, sub in boxes:
        _box(draw, (left, top, right, bottom), title, sub)
    for x in (340, 720, 1100, 1480):
        _arrow(draw, (x + 10, 267), (x + 65, 267))
    _box(draw, (650, 485, 1250, 650), "4 · USER EXPERIENCE", "React upload, progress, results")
    _arrow(draw, (1320, 355), (1190, 480)); _arrow(draw, (1670, 355), (1255, 540))
    diagrams["pipeline"] = _save(image, "pipeline_flow.png")

    image, draw = _new_canvas(1800, 790)
    _centre_text(draw, (0, 12, 1800, 85), "TWO HONEST OUTPUT MODES", _font(35, True), f"#{GREEN}")
    _box(draw, (690, 110, 1110, 230), "INPUT IMAGE")
    _arrow(draw, (800, 230), (430, 325)); _arrow(draw, (1000, 230), (1370, 325))
    _box(draw, (90, 325, 770, 485), "PNG / JPG / UNTRUSTED TIFF", "no trustworthy map grid")
    _box(draw, (1030, 325, 1710, 485), "VALID GEOREFERENCED GEOTIFF", "CRS + affine transform")
    _arrow(draw, (430, 485), (430, 575)); _arrow(draw, (1370, 485), (1370, 575))
    _box(draw, (90, 575, 770, 750), "RELATIVE SURFACE", "units: rel · visualization only", fill="F3EFE7")
    _box(draw, (1030, 575, 1710, 750), "ESTIMATED ABSOLUTE DSM", "only with SRTM and/or GCP", fill="E4EEE9")
    diagrams["modes"] = _save(image, "two_modes.png")

    image, draw = _new_canvas(1800, 780)
    _centre_text(draw, (70, 20, 760, 95), "RASTER PIXEL GRID", _font(34, True), f"#{GREEN}")
    origin_x, origin_y, cell = 100, 150, 90
    for r in range(4):
        for c in range(6):
            box = (origin_x+c*cell, origin_y+r*cell, origin_x+(c+1)*cell-4, origin_y+(r+1)*cell-4)
            draw.rectangle(box, fill="#EEF3F0", outline="#8BA69A", width=2)
    cx, cy = origin_x + 3*cell + cell//2, origin_y + 1*cell + cell//2
    draw.ellipse((cx-13, cy-13, cx+13, cy+13), fill=f"#{GREEN}")
    _centre_text(draw, (80, 540, 750, 620), "row = 1, col = 3", _font(27), f"#{MUTED}")
    _arrow(draw, (740, 330), (990, 290))
    _box(draw, (1000, 190, 1710, 355), "AFFINE TRANSFORM", "pixel centre → map x/y")
    _arrow(draw, (1355, 355), (1355, 500))
    _box(draw, (1000, 505, 1710, 670), "COORDINATE SYSTEM (CRS)", "map x/y → Earth location")
    diagrams["pixel"] = _save(image, "pixel_to_earth.png")

    image, draw = _new_canvas(1800, 800)
    _centre_text(draw, (0, 15, 1800, 90), "GLOBAL PASS + OVERLAPPING LOCAL TILES", _font(34, True), f"#{GREEN}")
    draw.rectangle((60, 130, 980, 700), fill="#F7F9F8", outline=f"#{GREEN}", width=4)
    tile_boxes = [(95, 220, 405, 480), (340, 220, 650, 480), (585, 220, 895, 480), (220, 390, 530, 650), (465, 390, 775, 650)]
    fills = ["AFC8BD", "E5EEE9", "C8D8D1", "F0F4F2", "D5E2DC"]
    for box, fill in zip(tile_boxes, fills):
        draw.rectangle(box, fill=f"#{fill}", outline=f"#{GREEN}", width=3)
    _arrow(draw, (995, 400), (1100, 245))
    _box(draw, (1110, 145, 1740, 290), "ALIGN TILE SCALE + SHIFT")
    _arrow(draw, (1425, 290), (1425, 345))
    _box(draw, (1110, 350, 1740, 495), "KEEP FINE STRUCTURAL DETAIL")
    _arrow(draw, (1425, 495), (1425, 550))
    _box(draw, (1110, 555, 1740, 700), "COSINE BLEND + SEAM SMOOTH")
    diagrams["tiles"] = _save(image, "tiled_inference.png")

    image, draw = _new_canvas(1800, 800)
    _centre_text(draw, (0, 15, 1800, 90), "CALIBRATION AND FUSION", _font(36, True), f"#{GREEN}")
    _box(draw, (50, 125, 520, 280), "SRTM", "coarse absolute baseline")
    _box(draw, (50, 330, 520, 485), "RELATIVE DEPTH", "fine estimated structure")
    _box(draw, (50, 535, 520, 690), "GCPS", "known height anchors")
    _box(draw, (700, 300, 1090, 500), "CALIBRATION", "fit scale + offset")
    _box(draw, (1320, 300, 1750, 500), "ESTIMATED DSM", "baseline + calibrated detail")
    for y in (200, 405, 610):
        _arrow(draw, (525, y), (695, 400))
    _arrow(draw, (1095, 400), (1315, 400))
    diagrams["fusion"] = _save(image, "calibration_fusion.png")

    image, draw = _new_canvas(1900, 880)
    actors = [(170, "BROWSER\nREACT"), (620, "FASTAPI\nPERSON 5"), (1100, "PIPELINE\nPERSONS 1–3"), (1660, "THREE.JS\nPERSON 6")]
    for x, label in actors:
        _centre_text(draw, (x-150, 20, x+150, 120), label, _font(28, True), f"#{GREEN}")
        for y in range(135, 825, 24):
            draw.line((x, y, x, min(y+12, 825)), fill="#C9D5D0", width=3)
    steps = [
        (170, 620, 170, "POST /api/process"),
        (620, 170, 265, "job_id + queued"),
        (620, 1100, 360, "run preprocessing / depth / calibration"),
        (170, 620, 475, "GET /api/status/{id}"),
        (620, 170, 570, "progress / completed"),
        (170, 620, 665, "GET /api/results/{id}"),
        (620, 1660, 775, "heightmap.json + texture + metadata"),
    ]
    for x1, x2, y, label in steps:
        _arrow(draw, (x1, y), (x2, y))
        _centre_text(draw, (min(x1,x2), y-45, max(x1,x2), y-5), label, _font(20), f"#{MUTED}")
    diagrams["api"] = _save(image, "api_sequence.png")

    image, draw = _new_canvas(1800, 850)
    _centre_text(draw, (0, 10, 1800, 90), "HEIGHT SAMPLES → VERTICES → TRIANGLES", _font(36, True), f"#{GREEN}")
    def project(x, z, y):
        return (900 + int((x-z)*95), 570 + int((x+z)*34) - int(y*210))
    grid = []
    for rz in range(-3, 4):
        row = []
        for cx in range(-4, 5):
            height = 1.0*math.exp(-0.22*((cx-1.0)**2+(rz+0.4)**2)) + 0.14*math.sin(cx)*math.cos(rz)
            row.append(project(cx, rz, height))
        grid.append(row)
    for row in grid:
        draw.line(row, fill=f"#{GREEN}", width=3)
    for col in range(len(grid[0])):
        draw.line([grid[row][col] for row in range(len(grid))], fill=f"#{DARK_GREEN}", width=3)
    for row in grid:
        for point in row:
            draw.ellipse((point[0]-5, point[1]-5, point[0]+5, point[1]+5), fill="#AFC8BD", outline=f"#{DARK_GREEN}")
    _centre_text(draw, (80, 665, 650, 790), "X / Z\ncolumns and rows\nor ground distance", _font(25), f"#{MUTED}")
    _centre_text(draw, (1200, 665, 1730, 790), "Y\nelevation or\nrelative relief", _font(25), f"#{MUTED}")
    diagrams["mesh"] = _save(image, "mesh_diagram.png")
    return diagrams


def add_role_summary(doc, role: str, mission: str, input_text: str, output_text: str, connects: str):
    add_table(
        doc,
        ["Role", "What this module owns"],
        [
            ["Mission", mission],
            ["Input", input_text],
            ["Output", output_text],
            ["Connection", connects],
        ],
        widths=[1.25, 5.15],
        font_size=8.5,
    )


def add_judge_mini(doc, pairs: list[tuple[str, str]]) -> None:
    rows = [[q, a] for q, a in pairs]
    add_table(doc, ["Likely judge question", "Strong, simple answer"], rows, widths=[2.3, 4.1], font_size=8.1)


def build_document() -> None:
    if sha256(REFERENCE) != REFERENCE_SHA256:
        raise RuntimeError("The retained template reference changed before authoring.")
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(REFERENCE, OUTPUT_DOCX)
    doc = Document(OUTPUT_DOCX)
    clear_body(doc)
    configure_styles(doc)
    rebuild_footer(doc)
    for section in doc.sections:
        section.top_margin = Inches(0.82)
        section.bottom_margin = Inches(0.78)
        section.left_margin = Inches(0.88)
        section.right_margin = Inches(0.88)
        section.header_distance = Inches(0.3)
        section.footer_distance = Inches(0.32)

    doc.core_properties.title = "DepthWizard — Team Technical Understanding & Judge Preparation Guide"
    doc.core_properties.subject = "Repository-verified beginner guide for the six-person DepthWizard hackathon team"
    doc.core_properties.author = "DepthWizard Team"
    doc.core_properties.keywords = "DepthWizard, monocular depth, GeoTIFF, DSM, SRTM, GCP, Three.js, FastAPI, React"
    diagrams = make_diagrams()

    # Cover
    add_eyebrow(doc, "Hackathon project · repository-verified edition")
    p = doc.add_paragraph(style="Title")
    p.paragraph_format.space_before = Pt(54)
    p.add_run("DepthWizard")
    p = doc.add_paragraph(style="Subtitle")
    p.add_run("Single-View Height Estimation and Interactive 3D Flythrough")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    run = p.add_run("TEAM TECHNICAL UNDERSTANDING & JUDGE PREPARATION GUIDE")
    run.font.name = "Georgia"; run.font.size = Pt(13); run.font.bold = True; run.font.color.rgb = RGBColor.from_string(GREEN)
    set_paragraph_border(p, color=GREEN, size="16", space="9")
    add_body(doc, "A beginner-friendly guide to the complete six-person pipeline: image ingestion, monocular relative depth, geospatial calibration, API integration, web experience, and GPU-accelerated 3D terrain.")
    doc.add_paragraph()
    add_table(
        doc,
        ["Edition", "Verified against", "Audience"],
        [["5 September 2026", "Current local repository and test suite", "All six teammates and hackathon judges"]],
        widths=[1.5, 2.65, 2.25],
        font_size=8.4,
    )
    add_note(doc, "Core promise", "DepthWizard turns one RGB image into an explorable height surface. It reports relative height when no trustworthy elevation reference exists and estimated metric elevation only after valid georeferencing plus SRTM and/or GCP calibration.")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.add_run("Hackathon prototype · not survey-grade photogrammetry · not a LiDAR replacement").italic = True
    add_page_break(doc)

    # Document control and use
    add_h1(doc, "How to use this guide", new_page=False)
    add_body(doc, "This is both a learning document and a judging-day revision pack. Read Chapters 1–4 once as a team, then study your own person chapter. Before the demo, every teammate should review the limitations, the end-to-end data contract, the judge Q&A, and their one-page cheat sheet.")
    add_table(
        doc,
        ["Reading path", "Best for", "Sections"],
        [
            ["15-minute team overview", "Understanding the product story", "Executive summary, Chapters 1–4, limitations"],
            ["Role deep dive", "Preparing your own technical explanation", "Your person chapter + your cheat sheet"],
            ["Demo operator", "Running and troubleshooting the live system", "Integration, demo runbook, startup troubleshooting"],
            ["Judge preparation", "Handling critical questions honestly", "Validation, limitations, 40 suggested answers"],
        ],
        widths=[1.55, 2.05, 2.8],
    )
    add_h2(doc, "Meaning of status labels")
    add_table(
        doc,
        ["Label", "Meaning in this document"],
        [
            ["Implemented", "The behavior exists in the current repository and is exercised by source code or tests."],
            ["Optional", "Supported only when the user provides an additional file or configuration."],
            ["Prototype limitation", "A known boundary that must be explained rather than hidden."],
            ["Future improvement", "A sensible next step; it is not claimed as current functionality."],
        ],
        widths=[1.7, 4.7],
    )
    add_h2(doc, "Words we use carefully")
    add_bullets(doc, [
        "Relative depth: a scene-dependent ordering or geometry signal with no physical unit.",
        "Relative height: the same unitless signal displayed as vertical relief; the viewer labels it rel.",
        "Estimated absolute DSM: a reference-calibrated surface in metres, still subject to model, alignment, datum, and reference error.",
        "Ground height: not automatically available. Building or canopy height above ground requires a trustworthy ground DTM and subtraction.",
    ])
    add_page_break(doc)

    # TOC
    add_h1(doc, "Contents", new_page=False)
    add_body(doc, "The final PDF contains a generated table of contents with page numbers and clickable heading bookmarks.")
    toc = doc.add_paragraph()
    add_toc_field(toc)
    add_page_break(doc)

    # Executive summary
    add_h1(doc, "Executive summary", new_page=False)
    add_body(doc, "Generating useful 3D elevation information normally depends on stereo imagery, LiDAR, photogrammetry, radar interferometry, or an existing elevation product. DepthWizard explores a faster, lower-input route: infer relative scene structure from one optical image, anchor it to real-world elevation when suitable references exist, and turn the result into an interactive 3D terrain that a user can inspect immediately.")
    add_picture(doc, diagrams["pipeline"], "Figure 1. Current repository architecture. Person 5 orchestrates the processing modules while Person 4 presents Person 6’s viewer.")
    add_h2(doc, "The two honest operating modes")
    add_picture(doc, diagrams["modes"], "Figure 2. A file extension alone does not decide the output mode; trustworthy metadata and calibration inputs do.")
    add_note(doc, "The one sentence everyone should remember", "A single image can provide useful relative structure, but it cannot by itself establish exact elevation above sea level. DepthWizard adds absolute reference only when geospatial metadata and SRTM/GCP evidence make that possible.", warning=True)
    add_h2(doc, "What is currently implemented")
    add_table(
        doc,
        ["Capability", "Current implementation", "Status"],
        [
            ["Input", "PNG, JPG/JPEG, TIFF/GeoTIFF; bounded 3072-pixel working grid for large scenes", "Implemented"],
            ["Relative depth", "Depth Anything V2 Base by default; CUDA when available; global plus overlapping 768 px tiles", "Implemented"],
            ["Absolute calibration", "Local SRTM/elevation raster and/or CSV/JSON GCPs; linear, inverse, robust-linear candidates", "Optional"],
            ["Geospatial output", "Float32 absolute_dsm.tif plus NPY, preview, metrics, and metadata for calibrated jobs", "Implemented"],
            ["Web application", "React frontend built with Vite; FastAPI backend with background jobs and polling", "Implemented"],
            ["3D", "Three.js/WebGL terrain, texture/elevation/wireframe modes, overview, free flight, auto path, point HUD", "Implemented"],
            ["Standalone scale", "Single-process local hackathon workflow; no durable worker queue or cloud autoscaling", "Prototype limitation"],
        ],
        widths=[1.25, 4.2, 0.95],
        font_size=7.9,
    )

    # Problem and motivation
    add_h1(doc, "1 · Problem definition and why it matters")
    add_h2(doc, "Polished hackathon problem statement")
    add_body(doc, "High-quality elevation and 3D surface data are valuable, but the usual acquisition routes—LiDAR, stereo photogrammetry, interferometric radar, or commercial DEM products—may be unavailable, expensive, slow, or operationally inconvenient. DepthWizard investigates whether a single RGB remote-sensing image, combined when possible with coarse elevation references and control points, can produce an approximate surface model that is immediately understandable through an interactive flythrough.")
    add_body(doc, "The goal is not to defeat the laws of geometry or replace survey instruments. The goal is to create a practical, transparent pipeline that gives useful relative structure from ordinary imagery, upgrades that structure to an estimated metric surface when valid references exist, and makes uncertainty visible to the user.")
    add_h2(doc, "Why the problem is hard")
    add_table(
        doc,
        ["Difficulty", "Plain-language explanation", "DepthWizard response"],
        [
            ["One view is ambiguous", "Different 3D scenes can create a similar 2D picture.", "Use a pretrained prior and label the result as an estimate."],
            ["No automatic metres", "Model values do not contain sea-level elevation or a universal scale.", "Fit scale/offset from SRTM or GCPs when available."],
            ["Domain gap", "Models trained on natural perspective images may misread nadir aerial scenes.", "Use a stronger relative model, tiles, cautious smoothing, and explicit limitations."],
            ["Geospatial grids differ", "SRTM and optical pixels may use different coordinates and resolutions.", "Reproject with CRS and affine transforms before comparing pixels."],
            ["3D can mislead", "A beautiful textured mesh can look correct even when geometry is wrong.", "Keep relative/metric labels, source metadata, and raw measurement values visible."],
        ],
        widths=[1.2, 2.5, 2.7],
        font_size=8.0,
    )
    add_h2(doc, "Responsible application areas")
    add_bullets(doc, [
        "Rapid visual familiarization with a remote area when richer 3D data is not immediately available.",
        "Educational demonstrations of DSMs, CRSs, monocular depth, and 3D terrain generation.",
        "Early-stage planning or communication for infrastructure and landform discussions.",
        "Disaster-response research prototypes for situational visualization—only after independent validation before operational use.",
        "Remote-sensing and defence research demonstrations where uncertainty and provenance are preserved.",
    ])
    add_note(doc, "Confidence, not hype", "The strongest pitch is that DepthWizard integrates six difficult layers into one transparent workflow. Its value is accessibility, speed, and inspectability—not a claim of survey-grade elevation from one image.")

    # Concepts
    add_h1(doc, "2 · Concepts every teammate should understand")
    add_h2(doc, "RGB image, raster, and height surface")
    add_body(doc, "An RGB image stores three brightness values per pixel: red, green, and blue. A raster is simply an organized grid of cells. It is not “one raster per pixel.” One raster has rows and columns, and each band stores a value at every cell. A height map is another raster: each cell stores one height-like value instead of a color.")
    add_table(
        doc,
        ["Data", "Value at one cell", "Units", "Typical use"],
        [
            ["RGB image", "Red, green, blue brightness", "Usually 0–255 after display normalization", "Appearance / texture"],
            ["Relative depth", "Learned near/far score", "Unitless", "Relative shape"],
            ["SRTM/DEM", "Reference elevation", "Usually metres for the selected product", "Broad absolute terrain"],
            ["DSM", "Top-surface elevation", "Metres when calibrated", "Terrain + structures/canopy surface"],
            ["nDSM", "DSM − reliable ground DTM", "Metres", "Height above ground"],
        ],
        widths=[1.15, 2.1, 1.5, 1.65],
        font_size=8.0,
    )
    add_h2(doc, "DEM, DTM, DSM, and nDSM")
    add_body(doc, "DEM is a broad family term for elevation grids. DTM usually means a bare-earth terrain surface. DSM means the upper visible or sensed surface, which may include roofs and vegetation. nDSM means normalized DSM: the height above a ground surface. Real products use these terms inconsistently, so the safest description of DepthWizard’s calibrated result is “estimated/fused DSM or height surface.”")
    add_code(doc, "nDSM = estimated DSM − trustworthy ground DTM")
    add_h2(doc, "Relative depth versus absolute elevation")
    add_body(doc, "Suppose two pixels receive model values 0.8 and 0.2. That may tell us that one appears closer to the camera, but it does not mean 0.8 m and 0.2 m, or 80 m and 20 m. Absolute elevation needs a reference with known height and a compatible coordinate system.")
    add_h2(doc, "Does the model know roads, buildings, and canopy?")
    add_body(doc, "Not as an explicit rule. The current model is geometric, not a semantic segmentation system. It may infer a roof or canopy as closer to a nadir camera because of learned shape, texture, boundaries, and context, but it is not guaranteed to obey “road = low, building = high.” Reliable class-specific height would need semantic fusion, remote-sensing training, and a ground model.")

    add_h1(doc, "3 · Geospatial concepts without the jargon")
    add_picture(doc, diagrams["pixel"], "Figure 3. A pixel location becomes a map coordinate through an affine transform, then gains Earth meaning through the CRS.")
    add_table(
        doc,
        ["Term", "Beginner explanation", "Why it matters here"],
        [
            ["CRS", "The rulebook that gives coordinates a meaning on Earth.", "Without it, x=500000 and y=3000000 are just numbers."],
            ["Geographic CRS", "Coordinates are angular latitude/longitude, commonly degrees.", "Degree spacing must not be treated directly as metres."],
            ["Projected CRS", "A map projection expresses location on a flat grid, often in metres.", "Useful for physical X/Y scale and slope."],
            ["Affine transform", "Six coefficients mapping pixel column/row to map x/y.", "Preserves the image footprint, origin, pixel size, and rotation."],
            ["Pixel resolution", "Ground spacing represented by a pixel, when metadata is valid.", "Controls horizontal scale in the viewer and raster alignment."],
            ["Bounding box", "The left, bottom, right, and top map extent.", "Checks whether reference rasters overlap."],
            ["NoData", "A value or mask meaning “no valid observation here.”", "Prevents borders and gaps from becoming terrain cliffs."],
            ["Vertical datum", "The reference surface from which height is measured.", "Different datums can add a systematic vertical offset."],
        ],
        widths=[1.25, 2.5, 2.65],
        font_size=7.9,
    )
    add_h2(doc, "GeoTIFF is more than a TIFF file")
    add_body(doc, "A GeoTIFF stores raster pixels plus geospatial metadata such as a CRS and affine transform. A .tif extension alone is not proof of correct georeferencing. The current Person 1 code transforms the image centre to longitude/latitude and checks whether it falls inside the declared CRS area of use. If the tag is inconsistent, it keeps the declared CRS for diagnosis but blocks metric calibration.")
    add_h2(doc, "Why ordinary image resizing is not reprojection")
    add_body(doc, "A visual resize changes array dimensions. Reprojection asks which Earth location each destination pixel represents, then samples the source at that location. Person 3 uses rasterio.warp.reproject for SRTM/reference alignment because cv2.resize or a generic image resize does not understand CRS, bounds, rotation, or pixel origin.")

    add_h1(doc, "4 · Intended architecture versus verified implementation")
    add_h2(doc, "The intended conceptual flow")
    add_code(doc, "Image → preprocessing → relative depth → optional calibration → height surface → API → textured 3D viewer → website")
    add_h2(doc, "What the repository actually does")
    add_body(doc, "The browser starts the work, so Person 5 is not merely a stage after Person 3. React sends a multipart upload to FastAPI. The backend creates a job and invokes Persons 1, 2, and optionally 3 in sequence. It then prepares or copies a bounded heightmap.json. React polls the job and lazily loads Person 6’s Three.js viewer on the results page.")
    add_picture(doc, diagrams["api"], "Figure 4. Technically accurate request and data sequence. The backend remains the orchestrator throughout processing.")
    add_table(
        doc,
        ["Concern", "Owner", "Important boundary"],
        [
            ["Image integrity and metadata", "Person 1", "Does not estimate depth or create a DSM."],
            ["Relative geometry", "Person 2", "Does not invent metres or read CRS."],
            ["Absolute reference and DSM", "Person 3", "Requires valid georeferencing plus SRTM/GCP."],
            ["User interaction", "Person 4", "Does not perform heavy ML/GIS work in the browser."],
            ["Job/API integration", "Person 5", "Coordinates modules; does not own their algorithms."],
            ["3D visualization and probing", "Person 6", "Rendering quality does not validate height accuracy."],
        ],
        widths=[2.0, 1.1, 3.3],
        font_size=8.0,
    )

    # Person 1
    add_h1(doc, "5 · Person 1 — Image loading, GeoTIFF metadata, and preprocessing")
    add_role_summary(doc, "Person 1", "Protect the image grid and geospatial meaning while creating a safe model input.", "PNG/JPG/JPEG/TIFF/GeoTIFF", "rgb_model.png, valid_mask.npy, metadata.json, preview.png; optional rgb_original.npy", "Feeds RGB to Person 2 and grid/CRS/mask metadata to Person 3 and Person 5.")
    add_h2(doc, "What this module does from scratch")
    add_numbered(doc, [
        "Validate that the path exists and the extension is supported.",
        "Read PNG/JPEG with Pillow or TIFF/GeoTIFF with Rasterio.",
        "Convert band-first raster data into height × width × 3 RGB order.",
        "Build a valid-pixel mask from NoData information and non-finite values.",
        "Detect edge-connected near-black scanner/film borders when a TIFF declares no NoData value.",
        "Validate that a claimed CRS and affine transform are geographically plausible.",
        "Decode very large imagery directly to a bounded working grid whose long edge is at most 3072 pixels.",
        "Normalize each display/model channel with valid-pixel 2nd and 98th percentiles, leaving the original numeric array unchanged.",
        "Write stable files and metadata that downstream modules can trust.",
    ])
    add_h2(doc, "Bands and array orientation")
    add_body(doc, "Rasterio reads bands × rows × columns. Neural-network image code normally expects rows × columns × channels. The module moves the band axis once. If a TIFF has one band, the band is repeated three times; if it has two, the second band is repeated as the third channel. This creates an RGB-shaped input, but it does not claim the data is true natural color.")
    add_code(doc, "Rasterio: (bands, rows, columns) → model image: (rows, columns, 3)")
    add_h2(doc, "How a pixel retains its Earth location")
    add_body(doc, "The output metadata records the original and working dimensions plus scale factors. For a bounded grid, downstream code uses a correspondingly scaled affine transform rather than pretending the small grid has the original pixel spacing. Integer row/column coordinates refer to pixel centres for round-trip conversions.")
    add_h2(doc, "Actual files")
    add_table(
        doc,
        ["File", "Purpose", "Who uses it"],
        [
            ["depthwizard_person1/main.py", "CLI orchestration and readable error handling", "Backend adapter / direct users"],
            ["src/image_loader.py", "Input detection and Pillow path", "Person 1"],
            ["src/geotiff.py", "Raster loading, masks, border detection, CRS validation", "Person 1"],
            ["src/preprocessing.py", "Robust per-channel normalization", "Person 1"],
            ["src/coordinates.py", "Pixel ↔ map and map ↔ lon/lat helpers", "Persons 1 and 3 concepts"],
            ["src/output.py", "Writes the stable handoff package", "Persons 2, 3, and 5"],
        ],
        widths=[2.25, 2.65, 1.5],
        font_size=7.8,
    )
    add_h2(doc, "Input and output contract")
    add_table(
        doc,
        ["Output", "Grid / type", "Meaning"],
        [
            ["rgb_model.png", "Working H × W × 3, uint8", "Aligned model texture/input; normalized only for inference/display"],
            ["valid_mask.npy", "Working H × W, bool", "True where pixels are valid scene content"],
            ["metadata.json", "JSON", "Dimensions, CRS, transform, bounds, resolution, NoData, scale mapping, warnings"],
            ["preview.png", "≤1200 px preview", "Human inspection only; never use for pixel-aligned depth"],
            ["rgb_original.npy", "Original grid and dtype", "Optional large preservation copy; backend skips it to reduce duplication"],
        ],
        widths=[1.4, 1.55, 3.45],
        font_size=8.0,
    )
    add_h2(doc, "Common failure cases")
    add_bullets(doc, [
        "CRS tag but impossible coordinates: mark non-georeferenced and warn; do not calibrate.",
        "Black border interpreted as deep terrain: mask only edge-connected near-black background, not all dark pixels.",
        "Huge compressed GeoTIFF expands in memory: decode to the bounded model grid and skip the duplicate original array in web jobs.",
        "Normalization makes original data unrecoverable: keep rgb_original.npy optional and record that rgb_model.png is a display/model transform.",
        "Depth returned flipped or transposed: keep row/column conventions explicit and test orientation.",
    ])
    add_judge_mini(doc, [
        ("Why use GeoTIFF?", "It carries the CRS and pixel-to-map transform needed to place estimates on Earth and align external elevation."),
        ("What happens for PNG/JPG?", "We still estimate relative structure, but no Earth coordinates or metres are invented."),
        ("What is an affine transform?", "Six coefficients that map image column/row to map x/y, including scale, origin, and possible rotation."),
        ("Why preserve metadata?", "If preprocessing loses the mapping, Person 3 cannot align SRTM or export a correctly located DSM."),
        ("Does every TIFF contain correct georeferencing?", "No. We require a meaningful transform and CRS and validate their geographic plausibility."),
    ])
    add_h2(doc, "Person 1 — 30-second explanation")
    add_body(doc, "I built the ingestion layer. It accepts ordinary images and GeoTIFFs, converts their bands into a consistent RGB grid, masks NoData and scanner borders, validates georeferencing, and creates a bounded normalized image for the depth model. I also preserve the dimensions, CRS, affine transform, bounds, resolution, and scale mapping in metadata so later calibration and 3D rendering do not lose the connection between a pixel and its Earth location.")

    # Person 2
    add_h1(doc, "6 · Person 2 — Monocular relative depth estimation")
    add_role_summary(doc, "Person 2", "Infer image-aligned relative geometry from one RGB view.", "rgb_model.png and optional valid_mask.npy", "relative_depth.npy, relative_depth_preview.png, depth_metadata.json, heightmap.json", "Receives the prepared grid from Person 1 and supplies a unitless surface to Person 3 and Person 6.")
    add_h2(doc, "What monocular depth means")
    add_body(doc, "Monocular means one camera view. A network learns statistical cues that often correlate with geometry: perspective convergence, relative size, texture density, occlusion boundaries, shadows, horizon position, and the typical shapes of objects. It predicts a value for each pixel, creating a depth map. Because one picture lacks triangulation, the result is not uniquely determined and must be treated as an informed estimate.")
    add_h2(doc, "The model actually used")
    add_body(doc, "The current default is the pretrained depth-anything/Depth-Anything-V2-Base-hf checkpoint loaded through Hugging Face Transformers. The repository also exposes Depth Anything V2 Small/Large and MiDaS DPT Hybrid options. Base is the local RTX default because it preserves more detail than Small while remaining practical on the tested 8 GB GPU.")
    add_note(doc, "Direction of the numbers", "The current raw output is recorded as relative inverse depth: a larger value normally means closer to the camera. For a near-nadir aerial view, closer often corresponds to a higher surface—but this is not a semantic or metric guarantee.")
    add_h2(doc, "Large-scene inference")
    add_picture(doc, diagrams["tiles"], "Figure 5. A scene-wide pass stabilizes the global trend; overlapping local passes recover higher-frequency structure.")
    add_numbered(doc, [
        "Choose CUDA when torch.cuda.is_available(), otherwise CPU. CUDA inference uses float16 autocast and enables TF32 matrix multiplication where supported.",
        "For an image no larger than 1024 pixels on either side, run a single prediction.",
        "For a larger working image, run one scene-wide pass and overlapping 768-pixel tiles with 128-pixel overlap.",
        "Linearly align each tile’s arbitrary scale and offset to the global prediction.",
        "Keep the tile’s fine residual instead of allowing local low-frequency bias to create visible rectangles.",
        "Blend tiles with cosine weights and smooth only a narrow stitched-detail residual at seams.",
        "Restore the depth map to the Person 1 working grid and apply the valid mask as NaN.",
    ])
    add_h2(doc, "Numerical output versus grayscale preview")
    add_body(doc, "relative_depth.npy is the scientific handoff and is not forced to 0–1. The PNG is for human viewing. Its contrast uses sampled 2nd/98th percentile limits, a modest histogram-equalization blend, and a locally detrended detail layer. Those operations improve visibility only; they do not change the NPY array used for calibration or point measurement.")
    sample_rgb = ROOT / "depthwizard_person5" / "runtime" / "jobs" / "job_58053d3fc8a3" / "person1" / "rgb_model.png"
    sample_depth = ROOT / "depthwizard_person5" / "runtime" / "jobs" / "job_58053d3fc8a3" / "person2" / "relative_depth_preview.png"
    if sample_rgb.exists() and sample_depth.exists():
        add_two_pictures(doc, sample_rgb, sample_depth, "Actual bounded RGB model input from a completed repository job.", "Actual display-only relative-depth preview from the same aligned grid.")
    add_h2(doc, "Why Person 3 is still necessary")
    add_body(doc, "A relative prediction has no stable zero and no universal scale. It cannot state that a roof is 142.3 m above a datum. Person 3 pairs depth values with known reference heights and fits an approximate conversion before producing a metric surface.")
    add_h2(doc, "Actual files")
    add_table(doc, ["File", "Responsibility"], [
        ["depthwizard_person2/config.py", "Model choices and default checkpoint"],
        ["src/depth_model.py", "Device selection and pretrained model loading"],
        ["src/inference.py", "Single-pass/tiled prediction, alignment, detail merge, blending"],
        ["src/postprocessing.py", "Resize, finite-value handling, and external mask application"],
        ["src/output.py", "NPY, metadata, robust grayscale preview, and bounded heightmap JSON"],
    ], widths=[2.3, 4.1])
    add_h2(doc, "Failure cases and responsible answers")
    add_bullets(doc, [
        "Nadir aerial imagery may lack the perspective cues common in ground-level datasets.",
        "Roofs and tree crowns can be confused; shadows may be interpreted as shape.",
        "Repeated textures, water, clouds, haze, and weak-texture surfaces can fail.",
        "A model can create plausible but false geometry; visual plausibility is not validation.",
        "Scene-wide scale can drift from image to image; calibration is required for metres.",
    ])
    add_judge_mini(doc, [
        ("How can AI estimate depth from one picture?", "It learns statistical visual cues from many examples, but the answer remains ambiguous and approximate."),
        ("Is it measuring height?", "No. The current backbone produces relative inverse-depth-like values, not elevation in metres."),
        ("Why Depth Anything V2 Base?", "It gives stronger detail than Small and fits the tested 8 GB RTX workflow; Small remains available for CPU."),
        ("Does it know a road is ground?", "No explicit semantic rule. It may infer geometry from appearance, but we do not guarantee class-specific height."),
        ("How do you reduce tile seams?", "We align local tiles to a global pass, merge only fine residuals, cosine-blend overlaps, and lightly smooth the stitched residual."),
    ])
    add_h2(doc, "Person 2 — 30-second explanation")
    add_body(doc, "I implemented monocular relative-depth inference using Depth Anything V2 Base. The model predicts an image-aligned unitless depth signal from one RGB view. For large aerial scenes, I combine a global pass with overlapping local tiles so the global trend stays consistent while roofs, roads, and canopy boundaries keep more detail. The numerical output remains raw relative depth; only the grayscale preview is contrast-enhanced. Person 3 then supplies the missing absolute scale.")

    # Person 3
    add_h1(doc, "7 · Person 3 — SRTM/GCP calibration and DSM generation")
    add_role_summary(doc, "Person 3", "Place relative depth on an approximate metric elevation scale and preserve geospatial alignment.", "Valid GeoTIFF metadata, relative_depth.npy, optional SRTM/elevation raster, optional GCP CSV/JSON, optional reference DSM", "absolute_dsm.tif/.npy, preview_dsm.png, metadata.json, metrics.json", "Consumes Persons 1–2 outputs; Person 5 serves products to Persons 4 and 6.")
    add_h2(doc, "Why SRTM is useful—but not enough")
    add_body(doc, "The Shuttle Radar Topography Mission (SRTM) is a radar-derived elevation product, not a model we train. Common global 1-arc-second data is roughly 30 m spacing near the equator and is typically supplied in geographic WGS84 with metre elevations referenced to EGM96 for the cited product. It is valuable as a broad absolute anchor, but one coarse cell cannot resolve every road, roof, or tree in a high-resolution optical image.")
    add_body(doc, "Resampling SRTM to a smaller optical pixel size does not create new detail. It only puts the available coarse information onto the target grid. DepthWizard therefore uses SRTM mainly for low-frequency elevation and uses the calibrated monocular signal for estimated high-frequency variation.")
    add_h2(doc, "What a GCP contributes")
    add_body(doc, "A Ground Control Point (GCP) is a location where we know an elevation and either the source-image row/column or WGS84 longitude/latitude. At least two distinct valid points are needed to fit both scale and offset. More well-distributed, accurate points improve stability; clustered or incorrect points can bias the entire scene.")
    add_code(doc, "Pixel GCP: name,row,col,elevation_m\nWGS84 GCP: name,longitude,latitude,elevation_m")
    add_h2(doc, "CRS-aware alignment")
    add_numbered(doc, [
        "Read the RGB GeoTIFF target grid: CRS, affine transform, width, height, bounds, resolution, and NoData.",
        "If Person 1 used a bounded depth grid, derive a scaled affine transform that preserves the original outer footprint.",
        "Reproject SRTM and any independent reference DSM into the target CRS, transform, width, and height with Rasterio.",
        "Check exact shape, CRS, and transform compatibility before comparing or fusing arrays.",
        "Transform longitude/latitude GCPs into the target CRS, then map coordinates to row/column.",
    ])
    add_h2(doc, "Calibration candidates")
    add_table(doc, ["Candidate", "Equation", "When it may help", "Caution"], [
        ["Linear", "H = aD + b", "Raw depth behaves approximately linearly with height", "Global assumption can miss local perspective effects"],
        ["Inverse", "H = a(1/D) + b", "Model output behaves more like inverse depth", "Zero/invalid D must be excluded"],
        ["Robust linear", "Huber fit", "A few reference pairs are outliers", "Robust is not immune to bad or clustered controls"],
    ], widths=[1.1, 1.25, 2.15, 1.9], font_size=7.9)
    add_body(doc, "With 4–50 samples the code computes leave-one-out diagnostics; larger SRTM sample sets use a deterministic holdout. Automatic selection prefers the candidate with the best available validation RMSE. These diagnostics measure fit to the supplied calibration pairs, not independent scene accuracy.")
    add_h2(doc, "Fusion")
    add_picture(doc, diagrams["fusion"], "Figure 6. Coarse absolute reference and fine estimated structure are combined only after calibration.")
    add_code(doc, "depth_smooth = Gaussian smooth(calibrated_depth)\ndepth_detail = calibrated_depth − depth_smooth\nfinal_dsm = smoothed_aligned_srtm + α × depth_detail")
    add_body(doc, "The alpha parameter controls detail strength and defaults to 1.0. If only GCPs are supplied, the code exports the GCP-calibrated depth estimate without an SRTM baseline. Raw relative depth is never simply added to SRTM.")
    add_h2(doc, "SRTM-only calibration")
    add_body(doc, "When no GCPs are provided, the code smooths the relative depth to approximately the effective SRTM scale, pairs those coarse values with aligned SRTM heights, and fits a global mapping. This is a practical heuristic, not proof that the fine monocular surface is correct.")
    add_h2(doc, "Output and validation")
    add_table(doc, ["Output", "Meaning"], [
        ["absolute_dsm.tif", "Float32 geospatial raster, metre values, working CRS/transform/footprint, NoData −9999"],
        ["absolute_dsm.npy", "Numeric array used by the heightmap converter"],
        ["preview_dsm.png", "Colored human preview with a metre-labelled colorbar"],
        ["metadata.json", "Calibration source, candidates, coefficients, grid, ranges, fusion, limitations"],
        ["metrics.json", "MAE, RMSE, Pearson r, and bias when an independent reference DSM is supplied"],
    ], widths=[1.75, 4.65])
    add_table(doc, ["Metric", "What it tells us", "What it does not prove"], [
        ["MAE", "Average absolute vertical error", "Whether a few extreme errors exist"],
        ["RMSE", "Penalizes large errors more strongly", "Why the errors occurred"],
        ["Pearson r", "Similarity of spatial variation", "Correct vertical offset or scale"],
        ["Bias", "Average prediction minus reference", "Local accuracy"],
    ], widths=[1.05, 2.55, 2.8], font_size=8.0)
    add_h2(doc, "Vertical datum caution")
    add_body(doc, "Two datasets can share horizontal coordinates and still disagree vertically because their zero surfaces differ. Ellipsoidal height and orthometric height are not automatically interchangeable. The prototype records assumptions but does not implement a full geoid/datum transformation pipeline; mismatched references may create systematic offsets.")
    add_h2(doc, "What happens when calibration is absent")
    add_body(doc, "The web pipeline does not fail ordinary images. Person 5 copies Person 2’s relative NPY, preview, and heightmap into the result contract and marks is_absolute_elevation=false, elevation_units=relative, and calibration_source=Absolute elevation unavailable. Person 3’s direct CLI correctly requires SRTM or GCPs because its sole purpose is absolute calibration.")
    add_judge_mini(doc, [
        ("Why not use SRTM alone?", "It anchors broad elevation but is too coarse to describe individual small structures in high-resolution imagery."),
        ("How is calibration performed?", "We align reference and depth grids, sample paired values, fit scale and offset with candidate regressions, then fuse calibrated detail with a coarse baseline."),
        ("Is the DSM accurate?", "It is an estimated/fused DSM. Accuracy is scene- and reference-dependent and must be reported against independent LiDAR or reference DSM."),
        ("Can this replace LiDAR?", "No. LiDAR directly measures geometry; our single-view result is an accessible prototype with much larger uncertainty."),
        ("What if SRTM and depth disagree?", "We preserve SRTM as the low-frequency anchor and limit monocular influence to calibrated detail; disagreement is a validation warning, not something to hide."),
    ])
    add_h2(doc, "Person 3 — 30-second explanation")
    add_body(doc, "I convert the unitless depth map into an estimated geospatial height surface. I align local SRTM or known GCP elevations to the image grid, fit the depth-to-height scale and offset using linear, inverse, and robust candidates, and combine SRTM’s coarse absolute baseline with the model’s estimated fine detail. I export a GeoTIFF DSM and record calibration, validation, and datum limitations. Without trustworthy reference data, the system stays honestly relative.")

    # Person 4
    add_h1(doc, "8 · Person 4 — Frontend website and user experience")
    add_role_summary(doc, "Person 4", "Make the scientific workflow understandable, controllable, and honest for a user.", "User-selected image plus optional SRTM/GCP files; API status/results JSON", "Upload request, progress UI, previews, metadata, downloads, and embedded 3D viewer", "Calls Person 5’s API and hosts Person 6’s viewer component on the results page.")
    add_h2(doc, "Actual stack")
    add_body(doc, "The frontend is a React application built with Vite. React Router provides the Home, Analyze, and Results pages. The Three.js viewer is lazy-loaded only on the results page so the main application does not pay the large renderer bundle cost up front.")
    add_h2(doc, "User flow")
    add_numbered(doc, [
        "Choose PNG, JPG/JPEG, TIF, or TIFF and see a local preview when the browser supports it.",
        "Optionally open Advanced calibration and attach a GCP CSV/JSON and/or SRTM TIFF/HGT. The UI rejects calibration files for a non-TIFF source.",
        "Submit a multipart POST request to /api/process.",
        "Show processing status while polling /api/status/{job_id}; retry transient polling failures with increasing delay.",
        "Navigate to /results/{job_id} when the job completes.",
        "Display relative-depth and DSM previews, metadata, warnings, downloadable files, and the 3D viewer.",
    ])
    add_h2(doc, "Communicating uncertainty")
    add_body(doc, "The interface changes units and explanatory text based on is_absolute_elevation. Calibrated results show metres and a validation caution. Relative results show rel and explicitly explain that the values are not measured metres or guaranteed semantic class heights. This is part of the technical design, not merely copywriting.")
    add_h2(doc, "Actual files")
    add_table(doc, ["File", "Responsibility"], [
        ["src/pages/HomePage.jsx", "Project story and five-step overview"],
        ["src/pages/AnalyzePage.jsx", "Upload validation, calibration inputs, job creation, polling"],
        ["src/pages/ResultsPage.jsx", "Result discovery, previews, downloads, lazy viewer"],
        ["src/services/realApi.js", "Multipart upload and GET calls to FastAPI"],
        ["src/components/MetadataPanel.jsx", "Mode-aware units, ranges, model/calibration facts, warnings"],
        ["src/components/TerrainViewer.jsx", "Lifecycle wrapper around Person 6’s createTerrainViewer"],
        ["src/styles/global.css", "Responsive visual system and states"],
    ], widths=[2.45, 3.95], font_size=7.8)
    add_h2(doc, "Why heavy computation stays out of the browser")
    add_body(doc, "The browser is excellent for interaction and WebGL rendering but not ideal for running a large PyTorch geospatial pipeline on arbitrary client devices. The backend centralizes Python, Rasterio, model weights, CUDA selection, files, and errors. The browser receives bounded previews and heightmap JSON, while full NPY/GeoTIFF products remain downloadable.")
    add_judge_mini(doc, [
        ("What does the frontend do?", "It validates user choices, uploads files, shows progress and errors, presents results, and embeds the interactive viewer."),
        ("How does it communicate with the backend?", "Through HTTP: one multipart POST to create a job and GET requests for status, results, and files."),
        ("How do you show uncertainty?", "We change units and warnings by mode: metres only for reference-calibrated runs, rel otherwise."),
        ("What happens if the backend is down?", "The API wrapper reports a clear reachability error instead of silently failing."),
        ("Why React/Vite?", "React organizes stateful upload/result components; Vite provides a fast development server and compact production build."),
    ])
    add_h2(doc, "Person 4 — 30-second explanation")
    add_body(doc, "I built the React/Vite user experience. Users upload an image, optionally attach SRTM or GCP calibration, see progress and clear errors, and then explore previews, metadata, downloads, and the 3D viewer. The frontend talks to FastAPI through a small HTTP service layer. Most importantly, it distinguishes relative results from calibrated metre results so the interface does not overstate what the model knows.")

    # Person 5
    add_h1(doc, "9 · Person 5 — Backend and API integration")
    add_role_summary(doc, "Person 5", "Turn six separate modules into one reliable request-to-result workflow.", "Multipart upload: image plus optional srtm and gcp", "Job ID/status JSON, result metadata and safe file URLs", "Invokes Persons 1–3, prepares Person 6 data, and serves Person 4.")
    add_h2(doc, "Backend and API basics")
    add_body(doc, "A backend is the server-side part of the application. An API is a documented set of requests and responses. HTTP POST usually sends or creates data; HTTP GET retrieves data. JSON is the small text format used for job state and result metadata. Large binary products are served as files rather than inserted into JSON.")
    add_h2(doc, "Actual stack and endpoints")
    add_body(doc, "The implementation uses FastAPI and Uvicorn. CORS allows common localhost frontend ports. FastAPI BackgroundTasks starts processing after returning the job ID, so a large model run does not keep the upload request open.")
    add_table(doc, ["Method", "Route", "Purpose"], [
        ["GET", "/health", "Confirm that the backend is alive"],
        ["POST", "/api/process", "Validate/save image and optional SRTM/GCP; queue a job"],
        ["GET", "/api/status/{job_id}", "Return public status, stage, progress, or safe failure message"],
        ["GET", "/api/results/{job_id}", "Return metadata and URLs after completion"],
        ["GET", "/api/files/{job_id}/{filename}", "Download/display an allowed generated file"],
    ], widths=[0.75, 2.25, 3.4], font_size=8.0)
    add_h2(doc, "Stage ordering and progress")
    add_code(doc, "15% preprocessing → 45% depth estimation → 75% calibration/fallback → 100% completed")
    add_body(doc, "pipeline_runner.py uses sys.executable and argument lists, never shell=True. Every stage has a 30-minute default timeout, required-output checks, captured stdout/stderr, and a readable label. A failure stops later stages and writes private diagnostics to the job’s status.json while exposing only a safe browser message.")
    add_h2(doc, "Data and file safety")
    add_bullets(doc, [
        "Random job IDs match job_[0-9a-f]{12}; each job has isolated input, person1, person2, person3, and results directories.",
        "The default image limit is 500 MB; uploads are streamed in 1 MB chunks instead of read entirely into HTTP memory.",
        "Extensions and file signatures are validated; empty or mismatched files are rejected.",
        "HGT files keep names such as N28E077.hgt because GDAL derives their location from the tile name.",
        "Plain-filename checks and fixed search directories prevent traversal through the file endpoint.",
        "Large arrays are hard-linked when possible, with copy fallback, to avoid unnecessary multi-gigabyte duplication.",
    ])
    add_h2(doc, "Relative versus calibrated branch")
    add_body(doc, "If no calibration file or configured reference exists, the backend does not call Person 3’s absolute CLI. It exposes Person 2’s aligned relative surface and writes a calibration_report.json that explicitly marks absolute elevation unavailable. If SRTM or GCP is supplied, it first confirms Person 1 accepted the georeferencing, then invokes Person 3 and converts absolute_dsm.npy to a bounded heightmap.json.")
    add_h2(doc, "Result contract")
    add_table(doc, ["Result field", "Consumer", "Meaning"], [
        ["depth_preview_url", "Person 4", "Human-readable relative-depth PNG"],
        ["dsm_preview_url", "Person 4", "Absolute preview or relative fallback preview"],
        ["dsm_download_url", "User", "Absolute DSM/relative NPY download"],
        ["heightmap_url", "Person 6", "Bounded JSON grid—not a raw NPY/TIFF"],
        ["texture_url", "Person 6", "Aligned bounded RGB PNG"],
        ["metadata_url", "Persons 4 and 6", "Grid, mode, calibration, units, transform, limits"],
    ], widths=[1.65, 1.15, 3.6], font_size=7.9)
    add_h2(doc, "Prototype scaling boundary")
    add_body(doc, "FastAPI BackgroundTasks and local job folders are suitable for a hackathon machine. A production city-scale service should use object storage, a durable queue, dedicated GPU workers, retry/idempotency policies, resource quotas, authentication, observability, and tiled geospatial output rather than one large synchronous local filesystem workflow.")
    add_judge_mini(doc, [
        ("What is an API?", "A defined set of HTTP requests and responses that lets the frontend use backend processing without knowing its internal code."),
        ("Does the backend implement the model?", "It orchestrates the owner modules through subprocess adapters; Person 2 owns inference and Person 3 owns calibration."),
        ("What happens when a stage fails?", "Later stages stop, the job becomes failed, private logs remain on disk, and the browser receives a safe actionable message."),
        ("Why background jobs?", "The upload returns quickly with a job ID, and the frontend can poll during long model inference."),
        ("How do modules exchange data?", "Through stable files—PNG, NPY, GeoTIFF, and JSON—inside an isolated job directory."),
    ])
    add_h2(doc, "Person 5 — 30-second explanation")
    add_body(doc, "I built the FastAPI integration layer that turns separate modules into one application. The backend validates and streams uploads, creates an isolated job, runs preprocessing, depth, and optional calibration in order, prepares the browser heightmap, and exposes status, metadata, previews, and downloads through safe endpoints. If a stage fails, the pipeline stops cleanly and returns a useful error. If calibration is absent, it deliberately returns a relative result instead of inventing metres.")

    # Person 6
    add_h1(doc, "10 · Person 6 — 3D terrain rendering and flythrough")
    add_role_summary(doc, "Person 6", "Convert an aligned height grid and RGB image into an explorable WebGL terrain without changing reported measurements.", "heightmap.json, rgb_model.png, metadata.json", "Interactive Three.js scene, point elevation/slope HUD, overview/free flight/auto path", "Embedded by Person 4; data and URLs supplied by Person 5.")
    add_h2(doc, "Height map to mesh")
    add_picture(doc, diagrams["mesh"], "Figure 7. Every sampled height becomes a vertex. Two triangles per valid grid cell form a continuous surface.")
    add_body(doc, "A mesh is a set of vertices connected into faces, usually triangles. For each height sample, X comes from the column, Z from the row, and Y from the height above a baseline. Neighboring vertices are connected into triangles because graphics hardware is optimized to render triangles and any planar surface patch can be represented by them.")
    add_code(doc, "grid sample (row, col, height) → vertex (X = col, Y = height, Z = row)")
    add_h2(doc, "Physical versus relative scaling")
    add_body(doc, "For a calibrated projected raster, horizontal spacing comes from pixel resolution and the CRS unit conversion; vertical values are metres. For EPSG:4326/CRS84, the code approximately converts longitude/latitude degree spacing to local metres at the scene centre. For a relative result, the unitless range is clipped to its 1st/99th percentiles for display and scaled to a conservative relief equal to about 3% of the scene width.")
    add_h2(doc, "Rendering quality features in the current code")
    add_bullets(doc, [
        "At most 512 terrain samples on the long edge (about 262,000 vertices maximum) with bilinear reduction to reduce aliasing.",
        "Triangles touching invalid source pixels are removed, preventing masked scan borders from becoming walls.",
        "Two conservative display-only smoothing passes for relative geometry; raw values remain in elevationHeights for the HUD.",
        "RGB texture with mipmapping, linear filtering, and up to 16× anisotropy depending on GPU capability.",
        "WebGL antialiasing, sRGB output, ACES filmic tone mapping, soft PCF shadows, and a 2048×2048 shadow map.",
        "Device pixel ratio capped at 2× and automatic resize handling for stable performance.",
        "GPU badge reads the WebGL renderer string and warns when an integrated or software renderer is selected.",
    ])
    add_h2(doc, "Texture and UV mapping")
    add_body(doc, "UV coordinates tell the graphics engine which RGB pixel belongs on each vertex. The convention is deterministic: height[row, col] corresponds to image pixel (col, row), and the top raster row remains the top texture row. Texture alignment requires the RGB and height surface to describe the same crop and orientation.")
    add_h2(doc, "Controls that actually exist")
    add_table(doc, ["Mode/control", "Behavior"], [
        ["Overview", "Orbit with drag, pan with right-drag, zoom with wheel"],
        ["Fly", "Click viewport for mouse look; W/S forward/back; A/D strafe; Q or Space up; E or Shift down"],
        ["Auto", "24-second closed Catmull–Rom path; pause/resume/reset"],
        ["Texture / Elevation / Wireframe", "Switch visual material without changing data"],
        ["Relief", "0.2×–2× for relative mode; 0.2×–5× for metric mode"],
        ["Speed / Look", "Adjust movement speed and pointer sensitivity"],
        ["Centre reticle", "Reports aimed value, slope, source pixel, calibration source, and map coordinate when available"],
    ], widths=[1.8, 4.6], font_size=8.0)
    add_h2(doc, "How point values are calculated")
    add_body(doc, "A ray from the camera through the centre reticle intersects the mesh. The viewer maps that point back to grid coordinates and bilinearly samples the unmodified numeric height field. For metric data, local rise over horizontal metre spacing gives a physical slope. For relative data, the HUD labels the slope visual because the vertical scale is arbitrary.")
    add_h2(doc, "GPU use")
    add_body(doc, "There are two separate GPU workloads. PyTorch/CUDA runs model inference on the backend. WebGL renders the mesh in the browser. The renderer requests high-performance mode, but that is only a browser hint; the included launcher starts an isolated Chromium-family session with a high-performance GPU preference, and the badge confirms whether NVIDIA/RTX is actually selected.")
    add_judge_mini(doc, [
        ("How do you make 3D from 2D?", "Each height sample becomes a vertex, adjacent vertices become triangles, and the original RGB image is mapped with UV coordinates."),
        ("Why downsample?", "A browser does not need one vertex for every source pixel; bounded bilinear sampling preserves shape while keeping interaction smooth."),
        ("Why do A/D move sideways?", "They add/subtract the camera’s right vector—translation—while mouse movement changes yaw/pitch—rotation."),
        ("Are HUD heights real?", "Only calibrated jobs report estimated metres; uncalibrated jobs report rel and visual slope."),
        ("Does smoothing alter measurements?", "No. Smoothing changes displayed vertex positions only; the HUD samples the original height field."),
    ])
    add_h2(doc, "Person 6 — 30-second explanation")
    add_body(doc, "I turn the heightmap and aligned RGB image into an interactive Three.js terrain. Height samples become mesh vertices, valid neighbours become triangles, and UV coordinates place the original image on the surface. I added overview, free flight, auto flythrough, texture/elevation/wireframe modes, relief controls, GPU diagnostics, and a centre probe for height, slope, pixel, and map coordinates. Relative smoothing improves appearance without changing the values we report.")

    # Integration
    add_h1(doc, "11 · Complete system integration")
    add_h2(doc, "End-to-end data flow")
    add_table(doc, ["Step", "Producer → consumer", "Artifact / request", "Invariant"], [
        ["1", "Person 4 → Person 5", "multipart image; optional srtm/gcp", "Supported type and valid signature"],
        ["2", "Person 5 → Person 1", "source path, output directory, 3072 max grid", "No shell interpolation"],
        ["3", "Person 1 → Person 2", "rgb_model.png + valid_mask.npy", "Same working H × W and orientation"],
        ["4", "Person 1/2 → Person 3", "metadata.json + relative_depth.npy + references", "Valid CRS and aligned/derived grid"],
        ["5a", "Person 3 → Person 5", "absolute_dsm.npy/.tif + metadata", "Metric only after reference calibration"],
        ["5b", "Person 2 → Person 5", "relative depth/heightmap fallback", "Units remain relative"],
        ["6", "Person 5 → Person 6", "heightmap.json + texture + metadata URL", "Bounded JSON, not raw NPY"],
        ["7", "Person 5 → Person 4", "results JSON + file URLs", "Mode and warnings preserved"],
    ], widths=[0.45, 1.45, 2.55, 1.95], font_size=7.45)
    add_h2(doc, "The non-georeferenced path")
    add_code(doc, "PNG/JPG → bounded RGB → relative inverse depth → relative heightmap.json → rel-labelled 3D terrain")
    add_body(doc, "No CRS, transform, map coordinate, or metre elevation is invented. The image can still produce a visually useful relative scene, but model ambiguity remains.")
    add_h2(doc, "The calibrated GeoTIFF path")
    add_code(doc, "Valid GeoTIFF + SRTM/GCP → aligned references → calibrated detail + coarse baseline → estimated absolute DSM → metre-labelled 3D terrain")
    add_body(doc, "The path is only taken if Person 1 accepts the georeferencing and the user/configuration supplies at least one calibration source.")
    add_h2(doc, "Why dimensions are a team-wide contract")
    add_body(doc, "If a teammate flips, crops, transposes, or independently resizes one array without recording the mapping, a roof in the RGB image may receive a road’s height. The repository therefore preserves working-grid orientation, stores model-to-original scale, derives a CRS-preserving target grid for calibration, and uses explicit source/target dimensions in viewer metadata.")
    add_h2(doc, "What is downloaded versus displayed")
    add_table(doc, ["Need", "Format", "Reason"], [
        ["ML/GIS numeric work", "NPY or GeoTIFF", "Retains float values and, for GeoTIFF, geospatial metadata"],
        ["Human 2D inspection", "PNG preview", "Small and universally displayable; not a numeric product"],
        ["Browser terrain", "heightmap.json", "Easy fetch/parse with bounded sample count and validity mask"],
        ["Texture", "rgb_model.png", "Browser-compatible and aligned to the working grid"],
        ["Provenance", "metadata.json / metrics.json", "Records model, mode, calibration, grid, ranges, warnings, validation"],
    ], widths=[1.35, 1.2, 3.85], font_size=8.0)

    add_h1(doc, "12 · Validation, evaluation, and what good evidence looks like")
    add_h2(doc, "Current automated verification")
    add_table(doc, ["Suite", "Passed", "What it checks"], [
        ["Person 1", "4", "PNG/TIFF/GeoTIFF loading, round-trip, border mask, invalid CRS, CLI outputs"],
        ["Person 2", "5", "Shape/orientation, finite output handling, mask behavior, tiled seam/detail behavior"],
        ["Person 3", "7", "Alignment, bounded grids, GCPs, SRTM fusion, metrics, invalid CRS rejection"],
        ["Person 5", "8", "Upload/API contracts, safety, stage results and failure behavior"],
        ["Person 6", "8", "Camera axes, JSON loader, physical extent, nodata triangles, sampling, display-only smoothing"],
        ["Frontend build", "Pass", "React/Vite production compilation; 48 modules transformed"],
    ], widths=[1.25, 0.7, 4.45], font_size=7.9)
    add_note(doc, "Verified state", "On 5 September 2026, all 32 automated tests passed and the production Vite build completed. Rasterio/Starlette deprecation warnings were present, but there were no test failures.")
    add_h2(doc, "Accuracy study the team should run")
    add_numbered(doc, [
        "Select independent reference DSM/LiDAR scenes that were not used as GCPs or calibration inputs.",
        "Stratify scenes: dense urban, sparse built-up, hilly, forested, barren, and water/cloud edge cases.",
        "Align prediction and reference to the same horizontal CRS, footprint, resolution, mask, and vertical datum.",
        "Report MAE, RMSE, Pearson r, bias, valid-pixel coverage, and runtime; do not report only one attractive example.",
        "Compare baselines: SRTM alone, monocular relative structure, GCP-only calibration, and fused SRTM + GCP.",
        "Perform ablations: Small vs Base model, single pass vs tiled, alpha values, number/distribution of GCPs.",
        "Show error maps and failure examples alongside best cases.",
    ])
    add_h2(doc, "How to validate the 3D experience")
    add_bullets(doc, [
        "Texture/height registration: identifiable roof corners and roads should align with relief boundaries.",
        "Navigation: overview, free flight, rise/fall, pointer lock, auto path, pause/resume/reset.",
        "Performance: interactive frame rate on target hardware and bounded memory for large scenes.",
        "Measurement integrity: point HUD equals the unmodified height field at known test pixels.",
        "Mode honesty: no metres, map coordinates, or metric slopes in the uncalibrated path.",
    ])

    add_h1(doc, "13 · Important limitations and assumptions")
    add_body(doc, "Limitations are not an admission that the project failed. They define where the result is useful, what evidence is still needed, and how a responsible team would improve it.")
    add_table(doc, ["Limitation", "Effect", "Responsible mitigation / future work"], [
        ["Monocular scale ambiguity", "Relative values vary by scene and lack metres", "Reference calibration; never label raw output as metric"],
        ["Plausible hallucination", "The model can invent incorrect geometry", "Independent reference validation and uncertainty maps"],
        ["Hidden surfaces", "One image cannot reconstruct unseen sides", "Multi-view imagery or LiDAR for complete geometry"],
        ["Canopy/building confusion", "Top surfaces and ground are not separated reliably", "Remote-sensing fine-tuning + semantic classes + ground DTM"],
        ["SRTM is coarse", "Individual roofs/roads are unresolved", "Use it as broad baseline, not detailed truth"],
        ["SRTM surface bias", "Radar response may differ from bare earth or optical top surface", "Use product-specific quality layers and better local DEM"],
        ["CRS mismatch", "Reference pixels align to the wrong location", "Validate CRS/transform/bounds and reproject geospatially"],
        ["Vertical datum mismatch", "Systematic height offset", "Transform datums/geoid models before evaluation"],
        ["GCP quality/distribution", "Bad anchors distort scale and offset", "Surveyed points, spatial coverage, robust fitting, holdout tests"],
        ["No position in PNG/JPG", "No map coordinates or absolute elevation", "Remain relative unless external georeferencing is explicitly supplied"],
        ["Texture can look convincing", "Users may trust visually attractive but wrong geometry", "Keep provenance, mode, warnings, and point evidence visible"],
        ["Large scenes", "High RAM, GPU memory, disk, and latency", "Bounded working grid, tiles, browser downsampling; production tiling/queue"],
        ["Prototype deployment", "Single local process and local job storage", "Durable workers, object storage, monitoring, authentication"],
    ], widths=[1.4, 2.05, 2.95], font_size=7.35)
    add_note(doc, "Do not claim", "“We recover exact building height from any single satellite image,” “SRTM is our trained AI,” “every GeoTIFF is valid,” “the 3D model proves accuracy,” or “this replaces LiDAR/photogrammetry.”", warning=True)
    add_h2(doc, "What the team can confidently claim")
    add_bullets(doc, [
        "A complete local pipeline from upload to relative or reference-calibrated height surface and interactive 3D visualization.",
        "Repository-tested handling of large imagery, masks, geospatial metadata, tiled inference, optional SRTM/GCP calibration, API jobs, and 3D navigation.",
        "Explicit separation of numerical data from display-only contrast, smoothing, and relief exaggeration.",
        "A practical platform for further remote-sensing fine-tuning and validation—not a finished measurement instrument.",
    ])

    add_h1(doc, "14 · Demo runbook and troubleshooting")
    add_h2(doc, "Fastest reliable startup on Windows")
    add_code(doc, "cd \"C:\\Users\\preet\\OneDrive\\Documents\\ChatGPT\\Hackathon 2026 Depth wizard\"\n.\\start_depthwizard.cmd")
    add_body(doc, "The .cmd launcher works even when PowerShell execution policy blocks .ps1 files. It starts separate backend and frontend terminals, waits for both HTTP services, and opens the high-performance GPU viewer. Manual .cmd launchers also exist: start_backend.cmd, start_frontend.cmd, and open_gpu_viewer.cmd.")
    add_h2(doc, "If using PowerShell scripts")
    add_code(doc, "Set-ExecutionPolicy -Scope Process -ExecutionPolicy RemoteSigned\n.\\start_backend.ps1\n# new terminal at the same repository root\n.\\start_frontend.ps1")
    add_h2(doc, "Pre-demo checklist")
    add_bullets(doc, [
        "Run test_all.ps1 (or the .cmd launcher’s documented checks) before the event; keep one known-good sample ready.",
        "Verify http://127.0.0.1:8000/health and http://127.0.0.1:5173 before uploading.",
        "For a metric demo, use a genuinely valid georeferenced GeoTIFF and an overlapping SRTM/GCP file.",
        "Confirm backend metadata says ML compute: CUDA and the viewer GPU badge says NVIDIA/RTX rather than Intel/software.",
        "Explain the mode before showing the terrain: relative (rel) or estimated absolute (m).",
        "Keep a short screen recording as a fallback, but demonstrate the live point HUD if the machine is stable.",
    ])
    add_h2(doc, "Common startup problems")
    add_table(doc, ["Symptom", "Likely cause", "Fix"], [
        [".ps1 cannot be loaded", "Execution policy", "Use start_depthwizard.cmd, or set process-scoped RemoteSigned"],
        ["node is not recognized", "Node not on PATH", "Use the root launcher, which tries pnpm/npm and the bundled runtime"],
        ["Site cannot be reached", "Frontend exited or port 5173 not ready", "Read its terminal; verify the process stayed running"],
        ["Backend unreachable", "Uvicorn failed or wrong directory/environment", "Open /health and inspect the backend terminal"],
        ["Absolute elevation unavailable", "No valid SRTM/GCP or georeferencing rejected", "Provide compatible reference files and inspect the warning"],
        ["Integrated GPU badge", "Browser ignored high-performance hint", "Use open_gpu_viewer.cmd/ps1 and re-check badge"],
        ["3D cliffs at borders", "Invalid mask missing or old job output", "Reprocess with current Person 1/6 code"],
        ["Texture does not align", "Grid/crop/orientation mismatch", "Trace Person 1 metadata and width/height through every handoff"],
    ], widths=[1.45, 2.15, 2.8], font_size=7.8)

    # 40 Q&A
    add_h1(doc, "15 · Possible judge questions and suggested answers")
    add_body(doc, "Use these as talking points, not a script to memorize word for word. The strongest answer states what works, what evidence supports it, and what remains uncertain.")
    questions = [
        ("1. What problem are you solving?", "We convert one optical image into an approximate relative or calibrated height surface and make it immediately explorable in 3D. The calibrated path uses valid georeferencing plus SRTM and/or GCPs."),
        ("2. Why use a single view instead of stereo?", "Single images are far more available and operationally simple. We accept the ambiguity and treat monocular output as a prior, not as direct survey measurement."),
        ("3. How can one image contain depth information?", "Perspective, occlusion, texture, size, shadows, and scene context contain statistical cues. A pretrained network learns those cues, but multiple 3D scenes can still explain the same image."),
        ("4. What model do you use?", "The default is depth-anything/Depth-Anything-V2-Base-hf through Transformers. Small, Large, and MiDaS DPT Hybrid are configurable alternatives."),
        ("5. Does that model output metres?", "No. Our selected checkpoint produces scene-dependent relative inverse-depth-like values. We add approximate metric scale only through external reference calibration."),
        ("6. What does larger depth mean?", "For the current raw representation, larger generally means closer to the camera. The direction is stored in depth_metadata.json."),
        ("7. Does the model know buildings are high and roads are low?", "Not as an explicit semantic rule. It infers geometric appearance, so class-specific behavior can fail. Semantic fusion is future work."),
        ("8. Why Depth Anything V2 Base?", "It offers stronger detail than Small while fitting the tested 8 GB RTX workflow. Base is a practical accuracy/performance compromise for the hackathon."),
        ("9. Why tiled inference?", "Shrinking a large aerial image to one small pass destroys roofs and boundaries. We combine a global prediction with overlapping local tiles to retain detail and scene consistency."),
        ("10. How do you prevent tile seams?", "Local tiles are scale/shift aligned to a global pass, reduced to fine residual detail, cosine blended across 128-pixel overlaps, and lightly seam-smoothed."),
        ("11. What is SRTM?", "A radar-derived elevation dataset from the Shuttle Radar Topography Mission. It is reference data, not an AI model we train."),
        ("12. What is SRTM’s spatial resolution?", "For the global 1-arc-second product, sample spacing is about 30 m near the equator. Exact spacing varies with latitude and product."),
        ("13. Why not just use SRTM?", "It gives broad absolute elevation but cannot resolve most individual buildings or narrow roads in a high-resolution optical scene."),
        ("14. What is a GCP?", "A known elevation anchor associated with a source pixel or geographic coordinate. We use valid points to estimate scale and offset."),
        ("15. How is calibration performed?", "We form reference height/depth pairs and compare linear, inverse, and Huber robust-linear mappings. Auto mode selects using cross-validated or holdout RMSE when available."),
        ("16. What is your fusion equation?", "We separate calibrated depth into coarse and fine components, then add alpha times its fine detail to a smoothed aligned SRTM baseline."),
        ("17. What is a DEM?", "A general elevation grid. Depending on the product, it may represent terrain or another surface; we always check the product definition."),
        ("18. What is a DSM?", "A digital surface model represents the upper sensed surface, which may include roofs and vegetation. Our output is an estimated/fused DSM."),
        ("19. What is a DTM or nDSM?", "A DTM is intended as bare earth. nDSM is DSM minus a reliable ground DTM, giving height above ground. Our prototype does not promise a reliable ground DTM."),
        ("20. What is a GeoTIFF?", "A TIFF raster with geospatial metadata such as a CRS and affine transform, allowing pixels to be placed on Earth."),
        ("21. What is a CRS?", "A coordinate reference system defines what map coordinate numbers mean and how they relate to Earth."),
        ("22. What is the affine transform?", "Six coefficients mapping pixel column and row to map x and y, including origin, resolution, and possible rotation."),
        ("23. Why is ordinary resize not enough for SRTM alignment?", "Resize knows only array dimensions. Reprojection uses both rasters’ CRS, transforms, and footprints so samples correspond to the same Earth locations."),
        ("24. What is a vertical datum?", "The zero reference for elevation. If SRTM, GCPs, and validation data use different vertical datums, their heights may differ by a systematic offset."),
        ("25. Is your DSM accurate?", "It is an estimate whose accuracy depends on scene, model, references, alignment, and datum. We only quantify accuracy against an independent reference DSM/LiDAR."),
        ("26. How would you validate it?", "Use held-out reference DSM/LiDAR, align grids and datums, mask invalid overlap, and report MAE, RMSE, correlation, bias, and error maps across scene types."),
        ("27. Can this replace LiDAR or photogrammetry?", "No. Those methods directly measure geometry with stronger constraints. DepthWizard is a fast visualization and research prototype for situations with limited inputs."),
        ("28. What happens for PNG/JPG?", "The pipeline still produces aligned relative depth and 3D relief, but units remain rel and no Earth coordinates or metre elevation are claimed."),
        ("29. What if a GeoTIFF has bad metadata?", "Person 1 checks that the CRS and affine transform are meaningful and geographically plausible. Bad georeferencing blocks absolute calibration."),
        ("30. How does the frontend communicate with the backend?", "React sends a multipart POST to create a job, polls GET status, then fetches GET results and file URLs."),
        ("31. Why FastAPI?", "It provides typed HTTP routes, file uploads, automatic API docs, background tasks, and clear Python integration with the ML/GIS modules."),
        ("32. What happens if a stage fails?", "The backend stops the pipeline, records the failed stage and private logs, marks the job failed, and returns a safe actionable message."),
        ("33. How do you make the 3D mesh?", "The bounded height grid becomes vertices; valid neighboring samples become triangles; UV coordinates map the aligned RGB image onto the surface."),
        ("34. Why triangles?", "GPUs render triangles efficiently, and a pair of triangles represents each grid cell without ambiguity."),
        ("35. Why downsample the terrain?", "One vertex per large source pixel would waste browser memory and frame time. A 512-sample long edge preserves overview shape while keeping interaction smooth."),
        ("36. Does vertical exaggeration change the data?", "No. It scales displayed vertex Y positions. The point HUD reads the original numeric field; relative slope is labelled visual."),
        ("37. Does the project use the GPU?", "Yes: PyTorch CUDA accelerates depth inference when available, and WebGL uses the browser-selected GPU for rendering. They are separate workloads."),
        ("38. How much GPU memory does it use?", "We tested the Base model on an 8 GB RTX 5060. Exact peak depends on model, tile size, and image; bounded 3072-pixel preprocessing and 768-pixel tiles control memory. We should profile before claiming a universal number."),
        ("39. What is the innovation?", "The contribution is the transparent integration: large-image relative inference, optional geospatial calibration, honest dual-mode outputs, stable API contracts, and real-time textured analysis in one local workflow."),
        ("40. What would you improve next?", "Fine-tune on aerial height datasets, add semantic/ground separation and uncertainty, automate high-quality DEM retrieval, reconcile vertical datums, validate on LiDAR benchmarks, and deploy durable GPU workers."),
    ]
    for index, (question, answer) in enumerate(questions):
        if index and index % 5 == 0:
            add_page_break(doc)
        add_h3(doc, question)
        add_body(doc, answer)

    # Cheat sheets
    cheats = [
        {
            "person": "Person 1",
            "title": "Image loading, GeoTIFF metadata, and preprocessing",
            "module": "Prepare a clean, aligned RGB grid while preserving or rejecting geospatial meaning correctly.",
            "input": "PNG/JPG/JPEG/TIFF/GeoTIFF.",
            "do": "Validate format; read bands; build masks; detect scan borders; validate CRS/transform; bounded decode; percentile-normalize; record grid mapping.",
            "output": "rgb_model.png, valid_mask.npy, metadata.json, preview.png; optional rgb_original.npy.",
            "why": "Every later module fails if orientation, validity, or pixel-to-Earth mapping is wrong.",
            "tech": "Python, Rasterio/GDAL, Pillow, NumPy, PyProj, Affine.",
            "concepts": "Raster grid; RGB bands; CRS; affine transform; NoData/mask.",
            "questions": "Why GeoTIFF? What happens for JPG? What is CRS? How does a pixel map to Earth? Why preserve metadata?",
            "short": "I own the ingestion layer. I create a consistent RGB model image, mask invalid regions, validate claimed georeferencing, and preserve the original-to-working grid mapping. Person 2 receives aligned RGB; Person 3 receives the CRS, transform, resolution, and mask. I never invent coordinates for PNG/JPG or for a TIFF whose metadata is geographically inconsistent.",
            "long": "My module is the foundation of the pipeline. A normal image contains pixel colors, while a GeoTIFF can also contain a CRS and affine transform. I read TIFFs with Rasterio and ordinary images with Pillow, convert band-first arrays into rows × columns × RGB, handle one/two-band fallbacks, and build a valid-pixel mask from NoData and non-finite values. For undeclared black scanner borders I mask only dark regions connected to the edge. Large images are decoded to a bounded 3072-pixel working grid, and the metadata records how that grid maps back to the source. Model RGB uses robust per-channel 2nd/98th percentile normalization, but the original numeric data is not silently overwritten. Most importantly, I validate whether georeferencing is actually plausible. Without a valid CRS and transform, the system remains relative. This protects Person 3’s alignment and prevents a visually plausible but geographically wrong DSM.",
        },
        {
            "person": "Person 2",
            "title": "Monocular relative depth estimation",
            "module": "Infer a unitless, image-aligned near/far surface from one RGB view.",
            "input": "rgb_model.png plus optional valid_mask.npy.",
            "do": "Load Depth Anything V2 Base; select CUDA/CPU; run single/global+tiled inference; align and blend tiles; validate/mask; create preview and JSON.",
            "output": "relative_depth.npy, depth_metadata.json, relative_depth_preview.png, heightmap.json.",
            "why": "Provides fine scene structure that a coarse DEM cannot resolve, while leaving absolute scale to reference calibration.",
            "tech": "PyTorch, CUDA, Transformers, Depth Anything V2, NumPy, Pillow.",
            "concepts": "Monocular ambiguity; relative inverse depth; learned cues; tiled inference; numerical vs display output.",
            "questions": "How can one image give depth? Is it metres? Which model? Does it know roads/buildings? What fails?",
            "short": "I run Depth Anything V2 Base to estimate relative inverse depth from one prepared RGB image. The result is aligned to the image but has no physical unit. For large scenes, a global pass establishes consistent structure and overlapping local tiles recover detail without obvious seams. The raw NPY remains unchanged; only the PNG preview is contrast-enhanced. Person 3 uses external references to add approximate metric scale.",
            "long": "A single image has no stereo baseline, so depth is fundamentally ambiguous. The pretrained model learns statistical cues such as perspective, texture, occlusion, size, and context; it produces a useful prior, not a measurement. Our default checkpoint is depth-anything/Depth-Anything-V2-Base-hf. It records larger raw values as closer and outputs relative inverse-depth-like values, not metres. I use CUDA with float16 autocast when available and CPU otherwise. If the working image is large, one scene-wide pass supplies the global trend and 768-pixel local tiles with 128-pixel overlap recover roofs, roads, and canopy boundaries. Each local prediction has its scale and offset aligned to the global map; only its fine residual is cosine blended, followed by very light seam smoothing. The output is restored to the Person 1 grid, invalid pixels become NaN, and metadata records the representation, model, device, shape, and tile count. Preview enhancement never changes the NPY used for calibration or measurement.",
        },
        {
            "person": "Person 3",
            "title": "SRTM/GCP calibration and DSM generation",
            "module": "Convert relative depth into an estimated metric geospatial surface when trustworthy references exist.",
            "input": "Valid GeoTIFF grid, relative depth, optional local SRTM/elevation raster, optional GCPs, optional reference DSM.",
            "do": "Reproject references; extract paired samples; fit linear/inverse/Huber mappings; fuse coarse baseline with calibrated detail; validate and export.",
            "output": "absolute_dsm.tif/.npy, preview_dsm.png, metadata.json, metrics.json.",
            "why": "Monocular depth has structure but no metres; SRTM/GCPs supply the missing absolute anchor.",
            "tech": "Rasterio, PyProj, NumPy, SciPy, scikit-learn Huber regression, Matplotlib.",
            "concepts": "DEM/DSM/nDSM; SRTM; GCP; reprojection; vertical datum.",
            "questions": "Why SRTM? Why not SRTM alone? What is calibration? How accurate is it? Can it replace LiDAR?",
            "short": "I align SRTM and/or known GCP elevations to the image grid, fit a depth-to-height scale and offset, and create an estimated DSM. SRTM provides coarse absolute elevation; calibrated monocular depth provides estimated fine structure. I export a georeferenced float DSM and record metrics and limitations. Without a valid reference, the website stays in relative mode instead of inventing metres.",
            "long": "Person 2 gives a scene-dependent relative signal, so my job is to add an approximate metric reference. I first read the GeoTIFF’s CRS, affine transform, footprint, and resolution. For a bounded depth grid I derive a scaled transform that preserves the original extent. SRTM or another elevation raster is reprojected with Rasterio into that exact grid; generic image resize would be wrong because it ignores Earth coordinates. GCPs can be source pixels with elevation or longitude/latitude with elevation. I build valid depth/height pairs and fit linear, inverse, and Huber robust-linear candidates, using cross-validation or a deterministic holdout where possible. If SRTM is present, I use its smoothed values as the low-frequency absolute baseline and add alpha times the calibrated high-frequency depth detail. The result is described as an estimated/fused DSM, not survey truth. Independent reference data can produce MAE, RMSE, correlation, and bias. I also warn about coarse SRTM, semantic errors, and vertical-datum mismatch.",
        },
        {
            "person": "Person 4",
            "title": "Frontend website and user experience",
            "module": "Guide the user from upload to understandable, mode-aware results and 3D exploration.",
            "input": "User files plus FastAPI status/results JSON and file URLs.",
            "do": "Validate selections; submit multipart job; poll progress; show errors/previews/metadata/downloads; lazy-load the viewer; label uncertainty.",
            "output": "Interactive React experience and API requests.",
            "why": "A technically strong pipeline is unusable—and easy to misinterpret—without clear states, units, warnings, and controls.",
            "tech": "React, React Router, Vite, JavaScript, CSS, Fetch/FormData.",
            "concepts": "Component state; HTTP; multipart upload; polling; relative vs metric UX.",
            "questions": "What does frontend own? How does it call backend? How are errors shown? Why lazy-load? How is uncertainty communicated?",
            "short": "I built the React/Vite interface. It accepts the image and optional calibration files, starts a backend job, polls progress, and displays previews, metadata, downloads, warnings, and the Three.js terrain. The UI explicitly switches between rel and metre modes so users do not confuse a relative model output with absolute elevation.",
            "long": "My responsibility is the full user journey, not only styling. On the Analyze page I accept PNG/JPG/GeoTIFF, validate the extension, allow optional GCP and SRTM uploads only with TIFF input, and send everything in FormData to FastAPI. The backend immediately returns a job ID; the page polls its status, retries a few transient failures, and navigates to the results page only after completion. Results resolve backend URLs for the relative-depth preview, DSM preview, downloads, metadata, texture, and heightmap. The Three.js viewer is lazy-loaded because its bundle is large and only needed on this page. A metadata panel shows dimensions, range, model, calibration source, and border masking. Calibrated outputs are labelled as estimated metres with a validation warning; uncalibrated outputs are rel with an explicit message that absolute elevation is unavailable. This makes the interface part of the scientific integrity of the project.",
        },
        {
            "person": "Person 5",
            "title": "Backend and API integration",
            "module": "Orchestrate files, modules, status, errors, and result discovery behind stable HTTP endpoints.",
            "input": "POSTed image plus optional SRTM/GCP files.",
            "do": "Stream/validate uploads; isolate jobs; run Persons 1–3 in order; choose calibrated/relative branch; convert heightmap; expose safe results.",
            "output": "Job/status/results JSON and generated-file downloads.",
            "why": "It turns separate scripts into one usable system and protects every data contract between them.",
            "tech": "Python, FastAPI, Uvicorn, multipart HTTP, subprocess, JSON, filesystem safety.",
            "concepts": "API; POST/GET; background job; adapter contract; failure propagation.",
            "questions": "What is an API? Why FastAPI? How are modules called? What if one fails? How does it scale?",
            "short": "I built the FastAPI bridge. It validates and streams uploads, creates a job, calls preprocessing, depth, and optional calibration in the correct order, prepares the browser heightmap, and exposes status, results, and downloads. Each failure stops safely and returns a clear message. Without calibration, the backend deliberately publishes relative output rather than a false metric DSM.",
            "long": "The frontend should not know how to run five different scripts, so my API owns orchestration. POST /api/process accepts image, srtm, and gcp multipart fields. Files are signature-checked, size-limited, streamed in 1 MB chunks, and stored inside a random isolated job. FastAPI returns the job ID and runs the pipeline as a background task. pipeline_runner.py invokes owner modules with sys.executable and argument lists, never a shell string, checks required outputs, captures logs, and updates public progress. Person 1 produces the bounded RGB/metadata, Person 2 produces depth, and Person 3 is called only for a valid georeferenced image with a reference. Otherwise I publish a labelled relative fallback. A converter creates bounded heightmap.json for Person 6. GET status, results, and file routes expose only safe information. The current local queue is appropriate for a hackathon; production would need durable workers, object storage, quotas, authentication, and monitoring.",
        },
        {
            "person": "Person 6",
            "title": "3D terrain rendering and interactive flythrough",
            "module": "Build an aligned textured terrain, camera system, and honest point-analysis interface in WebGL.",
            "input": "heightmap.json, rgb_model.png, metadata.json.",
            "do": "Load/validate grid; reduce to ≤512 samples; create vertices/triangles/UVs; remove invalid faces; render; navigate; probe raw values.",
            "output": "Overview, free flight, auto flythrough, material modes, relief controls, GPU badge, height/slope HUD.",
            "why": "The 3D layer turns a static raster into an intuitive surface that users can navigate and interrogate.",
            "tech": "Three.js 0.179.1, WebGL, Vite, OrbitControls, JavaScript.",
            "concepts": "Height map; mesh/triangles; UV texture; camera translation/rotation; display exaggeration vs data.",
            "questions": "How is mesh built? Why triangles/downsample? How is texture aligned? Are heights real? Does smoothing change values?",
            "short": "I convert the height grid into a Three.js mesh, map the aligned RGB image onto it, and provide overview, free-flight, and automatic camera modes. The viewer downsamples large grids for performance, removes invalid-border triangles, and uses high-quality texture filtering and soft shadows. A centre reticle reports the unmodified height, slope, source pixel, and map coordinate, with rel or metre labels depending on calibration.",
            "long": "Person 6 receives a compact JSON grid and aligned texture. Each row/column sample becomes a vertex, neighboring valid samples become triangles, and UV coordinates map the original image to the same orientation. I preserve physical X/Z aspect from pixel resolution; geographic degree grids receive an approximate local metre conversion. Metric DSM values use their elevation difference directly, while relative values are robustly clipped and conservatively scaled because they have no physical vertical unit. The renderer limits the long edge to 512 samples, uses bilinear reduction, discards triangles across invalid pixels, adds anisotropic mipmapped texture, ACES tone mapping, soft shadows, and a high-performance WebGL hint. Users can orbit, fly with W/S and A/D strafing, rise/fall, or follow an automatic path. The centre ray probes the original numeric field and calculates local slope; smoothing and relief affect display only. The GPU badge tells us whether the browser actually selected NVIDIA, Intel, or software rendering.",
        },
    ]
    for cheat in cheats:
        add_h1(doc, f"Cheat sheet · {cheat['person']} — {cheat['title']}")
        add_table(doc, ["Prompt", "My answer"], [
            ["My module", cheat["module"]],
            ["My input", cheat["input"]],
            ["What I do", cheat["do"]],
            ["My output", cheat["output"]],
            ["Why it matters", cheat["why"]],
            ["Main technologies", cheat["tech"]],
            ["5 concepts I must know", cheat["concepts"]],
            ["5 questions judges may ask", cheat["questions"]],
        ], widths=[1.55, 4.85], font_size=7.5)
        add_h2(doc, "30-second explanation")
        add_body(doc, cheat["short"])
        add_h2(doc, "2-minute detailed explanation")
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(cheat["long"])
        run.font.size = Pt(8.4)

    # Glossary and references
    add_h1(doc, "Appendix A · Glossary")
    glossary = [
        ("Absolute elevation", "Height referenced to a specified vertical datum; in this prototype, estimated only after calibration."),
        ("Affine transform", "Six coefficients that map raster pixel coordinates to map coordinates."),
        ("API", "A defined interface through which software components exchange requests and responses."),
        ("Bias", "Mean predicted elevation minus reference elevation."),
        ("Calibration", "Fitting a relationship between relative model values and known reference heights."),
        ("CRS", "Coordinate Reference System: rules that give map coordinates Earth meaning."),
        ("CUDA", "NVIDIA’s GPU-computing platform used here for PyTorch inference."),
        ("DEM", "General term for a raster of elevation values; product meaning must be checked."),
        ("DSM", "Digital Surface Model: elevation of the sensed upper surface, potentially including buildings and vegetation."),
        ("DTM", "Digital Terrain Model: usually a bare-earth surface."),
        ("GCP", "Ground Control Point with known location/elevation used as a calibration anchor."),
        ("GeoTIFF", "TIFF raster containing geospatial reference metadata."),
        ("Height map", "2D grid whose cells store height or relative-height samples for a mesh."),
        ("Inverse depth", "A representation where larger values often mean closer, opposite to ordinary distance."),
        ("MAE", "Mean Absolute Error: average magnitude of vertical errors."),
        ("Mesh", "Vertices connected into faces, here triangles, forming a 3D surface."),
        ("NoData", "A marker or mask identifying cells without valid observations."),
        ("nDSM", "Normalized DSM: surface elevation minus a reliable ground terrain model."),
        ("Orthometric height", "Height relative to a gravity-based geoid-like sea-level reference."),
        ("Raster", "A regular row-and-column grid of cells; bands store one or more values per cell."),
        ("Relative depth", "Scene-dependent near/far structure without a physical scale."),
        ("Reprojection", "Mapping raster samples between coordinate systems using geospatial metadata."),
        ("RMSE", "Root Mean Square Error: error metric that penalizes large deviations."),
        ("SRTM", "Shuttle Radar Topography Mission elevation data used as a coarse reference."),
        ("Texture / UV", "The image and coordinates that map its pixels onto a 3D surface."),
        ("Vertical datum", "Reference surface that defines the zero for elevation."),
        ("WebGL", "Browser graphics API used by Three.js to render on the GPU."),
    ]
    add_table(doc, ["Term", "Meaning"], [[a, b] for a, b in glossary], widths=[1.45, 4.95], font_size=7.7)

    add_h1(doc, "Appendix B · Verified repository file map")
    add_table(doc, ["Person", "Main location", "Key files"], [
        ["1", "depthwizard_person1", "main.py; config.py; src/geotiff.py, preprocessing.py, coordinates.py, output.py"],
        ["2", "depthwizard_person2", "config.py; src/depth_model.py, inference.py, postprocessing.py, output.py"],
        ["3", "depthwizard_person3", "main.py; src/alignment.py, calibration.py, fusion.py, validation.py, export.py; examples/"],
        ["4", "depthwizard_person4", "src/pages/*.jsx; src/components/*.jsx; src/services/realApi.js; src/styles/global.css"],
        ["5", "depthwizard_person5", "main.py; api/routes.py; config.py; file_manager.py; pipeline_runner.py"],
        ["6", "depthwizard_person6", "src/main.js, terrain.js, dataLoader.js, firstPerson.js, flythrough.js; tools/prepare_heightmap.py"],
        ["Shared", "repository root", "README.md; start_depthwizard.cmd; start_backend.*; start_frontend.*; open_gpu_viewer.*; test_all.ps1"],
    ], widths=[0.65, 1.8, 3.95], font_size=7.6)
    add_h2(doc, "Default configuration snapshot")
    add_table(doc, ["Setting", "Current default"], [
        ["Depth model", "depth_anything_v2_base / depth-anything/Depth-Anything-V2-Base-hf"],
        ["Model working-grid long edge", "3072 pixels"],
        ["Local inference tile / overlap", "768 / 128 pixels"],
        ["Browser terrain long edge", "512 samples"],
        ["Upload limit", "500 MB"],
        ["Pipeline stage timeout", "1800 seconds"],
        ["Frontend / backend", "React + Vite / FastAPI + Uvicorn"],
        ["3D library", "Three.js 0.179.1"],
    ], widths=[2.15, 4.25], font_size=8.0)

    add_h1(doc, "Appendix C · References and further reading")
    add_body(doc, "The repository source code and test suite are the authority for implementation details in this guide. The following primary or official sources support model, geospatial, elevation, API, and rendering concepts.")
    references = [
        ("[1] Depth Anything V2 paper", "Yang et al., arXiv:2406.09414", "https://arxiv.org/abs/2406.09414"),
        ("[2] Depth Anything V2 Base model card", "Hugging Face model repository and usage notes", "https://huggingface.co/depth-anything/Depth-Anything-V2-Base-hf"),
        ("[3] SRTM product specifications", "USGS EROS: WGS84, EGM96, metre units, 1 arc-second ≈30 m", "https://www.usgs.gov/centers/eros/science/usgs-eros-archive-digital-elevation-shuttle-radar-topography-mission-srtm"),
        ("[4] SRTM Collection User Guide", "NASA LP DAAC/USGS Version 3 guide", "https://lpdaac.usgs.gov/documents/179/SRTM_User_Guide_V3.pdf"),
        ("[5] Rasterio reprojection", "Official Rasterio documentation", "https://rasterio.readthedocs.io/en/stable/topics/reproject.html"),
        ("[6] FastAPI background tasks", "Official FastAPI tutorial", "https://fastapi.tiangolo.com/tutorial/background-tasks/"),
        ("[7] Three.js WebGLRenderer", "Official renderer options, GPU capability, tone mapping, shadows", "https://threejs.org/docs/pages/WebGLRenderer.html"),
    ]
    add_table(doc, ["Source", "Use in this guide", "URL"], [[a, b, c] for a, b, c in references], widths=[1.65, 2.7, 2.05], font_size=7.3)
    add_h2(doc, "Final team message")
    add_body(doc, "DepthWizard is strongest when the team presents it as an integrated, evidence-aware prototype: relative geometry from a single image; metric anchoring only from trustworthy references; geospatial metadata preserved; numerical and visual transformations separated; and uncertainty explained directly. That honesty makes the engineering more credible, not less ambitious.")

    # Package settings and save.
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")
    compat = settings.find(qn("w:compat"))
    if compat is None:
        compat = OxmlElement("w:compat")
        settings.append(compat)
    doc.save(OUTPUT_DOCX)
    if sha256(REFERENCE) != REFERENCE_SHA256:
        raise RuntimeError("The retained template reference changed during authoring.")
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    build_document()
